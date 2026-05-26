"""신규 레퍼런스를 Claude가 '의미상 맞는 문장에만' 인용 삽입 — 임베딩 클럼핑 대신 판단 기반.

절차:
  1) PubMed 신규 후보 수집(초록 포함)
  2) Introduction/Discussion에 대해 Claude에게 {{PMID}} 태그를 '지지하는 문장 뒤에만' 삽입 요청
     (관련 없으면 사용 금지, 기존 [n]·숫자·문장 보존)
  3) 등장 순서로 29번부터 번호 부여, 사용된 것만 참고문헌에 추가
  4) Word/EndNote/BibTeX 재빌드
실행: python scripts/integrate_citations_llm.py <restyled.md> <out_basename>
"""
from __future__ import annotations
import io, re, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

QUERIES = [
    "artificial sweetener depression adolescents",
    "non-nutritive sweetener mental health youth",
    "diet beverage depression cohort",
    "sugar-sweetened beverage adolescent depression",
    "ultra-processed food adolescent depression mental health",
    "soft drink depression gut microbiome",
]


def main():
    import warnings; warnings.filterwarnings("ignore")
    from src.config.env import bootstrap; bootstrap()
    from src.export.reference_library import (
        search_pubmed, _fetch_pubmed_xml, _parse_pubmed_xml, format_vancouver,
        to_endnote_xml, to_bibtex,
    )
    from src.export.citation_workflow import _parse_vancouver_entry
    from src.llm import get_llm_client

    src = Path(sys.argv[1] if len(sys.argv) > 1 else "data/exports/ZCB_yoosun_claude.md")
    base = sys.argv[2] if len(sys.argv) > 2 else "ZCB_final_yoosun"
    text = src.read_text(encoding="utf-8")

    mt = re.search(r"(?m)^#\s+(.+?)\s*$", text); title = mt.group(1).strip()
    mref = re.search(r"(?im)^##\s*references?\s*$", text)
    body, head = text[:mref.start()], None
    ref_lines = [l.rstrip() for l in text[mref.end():].splitlines() if l.strip()]
    n_existing = len(ref_lines)

    parts = re.split(r"(?m)^##\s+(.+?)\s*$", body)
    head = parts[0].rstrip()
    secs = [[parts[i].strip(), (parts[i + 1] or "").rstrip()] for i in range(1, len(parts) - 1, 2)]
    sec_by = {s[0].lower(): idx for idx, s in enumerate(secs)}

    # 1) 후보 수집
    existing_lc = " ".join(ref_lines).lower()
    cand, seen = {}, set()
    for q in QUERIES:
        try:
            for r in _parse_pubmed_xml(_fetch_pubmed_xml(search_pubmed(q, max_results=4))):
                if not r.pmid or r.pmid in seen or not r.abstract:
                    continue
                seen.add(r.pmid)
                if (r.title or "").lower()[:40] in existing_lc:
                    continue
                cand[r.pmid] = r
        except Exception as e:
            print(f"  [query 실패] {q}: {str(e)[:70]}")
    print(f"PubMed 신규 후보(초록 보유) {len(cand)}개")

    cand_block = "\n".join(
        f"{{{{{r.pmid}}}}}  {(r.authors[0] if r.authors else '?')} et al. {r.title}. "
        f"{r.journal} {r.year}.\n   ABSTRACT: {(r.abstract or '')[:550]}"
        for r in cand.values()
    )

    client = get_llm_client(provider="anthropic", task="paper_writing")

    def integrate(section_name, section_text):
        sysp = (
            "You are inserting citations into a finished medical paper section. "
            "You are given candidate references, each tagged {{PMID}}. "
            "Insert a reference's {{PMID}} tag IMMEDIATELY after the specific sentence whose claim it "
            "directly supports — and ONLY if it genuinely supports that claim. It is correct to use "
            "FEWER references than offered; skip any that do not fit this section. Multiple tags may "
            "follow one sentence. ABSOLUTE RULES: do not alter any wording, do not touch existing "
            "[n] markers or any number/statistic, only insert {{PMID}} tags. Return the full section "
            "text with tags inserted, nothing else."
        )
        prompt = (f"CANDIDATE REFERENCES:\n{cand_block}\n\n"
                  f"SECTION ({section_name}) — insert {{{{PMID}}}} tags at supporting sentences:\n\n{section_text}")
        return client.generate(prompt, system_prompt=sysp, task="paper_writing", max_tokens=4000)

    for nm in ("introduction", "discussion"):
        if nm in sec_by:
            si = sec_by[nm]
            try:
                out = integrate(nm, secs[si][1])
                # 안전: 기존 [n] 마커 수가 줄지 않았는지 확인
                if out and out.count("[") >= secs[si][1].count("["):
                    secs[si][1] = out.strip()
                    print(f"[OK] {nm}: 태그 삽입 (provider={getattr(client,'model','?')})")
                else:
                    print(f"[KEEP] {nm}: 출력 이상 → 원문 유지")
            except Exception as e:
                print(f"[KEEP] {nm}: 실패 {str(e)[:80]}")

    # 3) 등장 순서로 번호 부여 (29부터), 사용된 PMID만
    full = "\n\n".join(c for _, c in secs)
    order = []
    for m in re.finditer(r"\{\{(\d+)\}\}", full):
        if m.group(1) in cand and m.group(1) not in order:
            order.append(m.group(1))
    pmid2num = {pmid: n_existing + i + 1 for i, pmid in enumerate(order)}
    print(f"\n실제 사용된 신규 레퍼런스: {len(order)}개 → [{n_existing+1}..{n_existing+len(order)}]")

    # 태그 → [n] 치환 (연속 태그는 묶기)
    def repl(s):
        # 인접한 }}{{ 사이 공백 정리 후 그룹화
        s = re.sub(r"\}\}\s*\{\{", "}}{{", s)
        def grp(mt):
            nums = sorted(int(pmid2num[p]) for p in re.findall(r"\{\{(\d+)\}\}", mt.group(0)) if p in pmid2num)
            return "[" + ", ".join(map(str, nums)) + "]" if nums else ""
        s = re.sub(r"(?:\{\{\d+\}\})+", grp, s)
        return s
    for s in secs:
        s[1] = repl(s[1])

    # 참고문헌: 기존 + 사용된 신규(번호순)
    for pmid in order:
        i = pmid2num[pmid]
        ref_lines.append(format_vancouver(cand[pmid], i))

    # 4) 저장 + 빌드
    new_md = head + "\n\n" + "\n\n".join(f"## {l}\n{c}" for l, c in secs) + \
        "\n\n## References\n" + "\n".join(ref_lines) + "\n"
    Path(f"data/exports/{base}.md").write_text(new_md, encoding="utf-8")

    from docx import Document
    doc = Document(); doc.add_heading(title, 0)
    lab = {"abstract": "Abstract", "introduction": "Introduction", "methods": "Methods",
           "results": "Results", "discussion": "Discussion"}
    for l, c in secs:
        doc.add_heading(lab.get(l.lower(), l), level=1)
        for para in c.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    doc.add_heading("References", level=1)
    for line in ref_lines:
        doc.add_paragraph(line)
    b = io.BytesIO(); doc.save(b); Path(f"data/exports/{base}.docx").write_bytes(b.getvalue())

    all_refs = [_parse_vancouver_entry(re.sub(r"^\s*\d+[.)]\s*", "", l)) for l in ref_lines]
    Path(f"data/exports/{base}.xml").write_bytes(to_endnote_xml(all_refs).encode("utf-8"))
    Path(f"data/exports/{base}.bib").write_bytes(to_bibtex(all_refs).encode("utf-8"))

    print("\n신규 인용 (배치 순):")
    for pmid in order:
        r = cand[pmid]
        print(f"  [{pmid2num[pmid]}] {(r.authors[0] if r.authors else '?')} et al. {r.title[:78]} ({r.journal} {r.year})")
    print(f"\n참고문헌 총 {len(ref_lines)}개 (기존 {n_existing} + 신규 {len(order)})")
    print(f"출력: {base}.md/.docx/.xml/.bib → data/exports/")


if __name__ == "__main__":
    main()
