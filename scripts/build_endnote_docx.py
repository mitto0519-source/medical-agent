"""ZCB_FINAL.md → EndNote CWYW 라이브 필드(EN.CITE travelling library 임베드) docx 생성.

본문 [n]/[n,m]/[n-m,...] 마커를 Word 필드 코드(begin/instrText/separate/displayText/end)로 변환.
각 필드에 EN.CITE XML(<record> 포함) 임베드 → EndNote에서 '라이브 인용'으로 인식, 별도 라이브러리 import 없이 Format Bibliography 가능.

한계(정직): 실제 EndNote가 깔린 환경에서 마지막 검증은 사용자 측. 필드 XML 구조 well-formed/Word 호환은 docx 파싱으로 자동검증함.

실행: python scripts/build_endnote_docx.py [in.md] [out_base]
"""
from __future__ import annotations
import io, re, sys
from html import escape as xml_escape
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.export.citation_workflow import _parse_vancouver_entry
from src.export.reference_library import format_vancouver


# ── EndNote XML 빌더 ──────────────────────────────────────────────────────────

def _esc(s: str) -> str:
    return xml_escape(s or "", quote=False)


def _record_xml(ref, rec_num: int) -> str:
    """단일 Reference → EndNote <record> XML (travelling library용, 한 줄)."""
    authors_xml = "".join(
        f'<author><style face="normal" font="default">{_esc(a)}</style></author>'
        for a in (ref.authors or [])
    )
    return (
        "<record>"
        f'<source-app name="EndNote" version="20"/>'
        f"<rec-number>{rec_num}</rec-number>"
        f'<foreign-keys><key app="EN" db-id="medagent" key-type="0">{rec_num}</key></foreign-keys>'
        f'<ref-type name="Journal Article">17</ref-type>'
        f"<contributors><authors>{authors_xml}</authors></contributors>"
        f'<titles>'
        f'<title><style face="normal" font="default">{_esc(ref.title)}</style></title>'
        f'<secondary-title><style face="normal" font="default">{_esc(ref.journal)}</style></secondary-title>'
        f'</titles>'
        f'<periodical><full-title><style face="normal" font="default">{_esc(ref.journal)}</style></full-title></periodical>'
        f'<pages><style face="normal" font="default">{_esc(ref.pages)}</style></pages>'
        f'<volume><style face="normal" font="default">{_esc(ref.volume)}</style></volume>'
        f'<number><style face="normal" font="default">{_esc(ref.issue)}</style></number>'
        f'<dates><year><style face="normal" font="default">{_esc(ref.year)}</style></year></dates>'
        f'<electronic-resource-num><style face="normal" font="default">{_esc(ref.doi)}</style></electronic-resource-num>'
        f'<accession-num><style face="normal" font="default">{_esc(ref.pmid)}</style></accession-num>'
        "</record>"
    )


def _build_encite(numbers: list, refs_by_num: dict, display: str) -> str:
    """본문 [...] 한 개 → ADDIN EN.CITE ... instrText 문자열 (raw < > 포함; instrText로 자동 escape됨)."""
    cites = []
    for n in numbers:
        ref = refs_by_num.get(n)
        if ref is None:
            continue
        # Vancouver "Lastname FM" → split()[0]이 surname (이전엔 [-1]로 이니셜만 잡혔던 버그)
        last = (ref.authors[0].split()[0] if ref.authors else "Anon")
        year = ref.year or "n.d."
        cites.append(
            f"<Cite>"
            f"<Author>{_esc(last)}</Author>"
            f"<Year>{_esc(year)}</Year>"
            f"<RecNum>{n}</RecNum>"
            f"{_record_xml(ref, n)}"
            f"</Cite>"
        )
    return (
        f" ADDIN EN.CITE <EndNote>"
        + "".join(cites)
        + f"<DisplayText>{_esc(display)}</DisplayText>"
        + f"</EndNote>"
    )


# ── 본문 [n] 토큰 파서 ─────────────────────────────────────────────────────────

_BRACKET = re.compile(r"\[([\d,\s\-]+)\]")


def _expand(bracket_inner: str) -> list:
    """'5-7, 15-17, 32' → [5,6,7,15,16,17,32] (정렬·중복제거)."""
    out = set()
    for tok in bracket_inner.split(","):
        tok = tok.strip()
        if "-" in tok:
            a, b = tok.split("-", 1)
            try:
                out.update(range(int(a), int(b) + 1))
            except ValueError:
                continue
        elif tok.isdigit():
            out.add(int(tok))
    return sorted(out)


# ── Word 필드 삽입 ────────────────────────────────────────────────────────────

def _add_field(paragraph, instr_text: str, display_text: str):
    """문단에 begin/instrText/separate/displayText/end 5-run 필드 시퀀스 추가."""
    # begin
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    r._r.append(fc)
    # instrText (lxml이 < > 자동 escape)
    r = paragraph.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = instr_text
    r._r.append(it)
    # separate
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate")
    r._r.append(fc)
    # display
    paragraph.add_run(display_text)
    # end
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end")
    r._r.append(fc)


def _emit_paragraph(doc, text: str, refs_by_num: dict, style=None):
    """텍스트 한 문단 → [n] 토큰을 EN.CITE 필드로, 나머지는 일반 run으로 분해해 add."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    pos = 0
    for m in _BRACKET.finditer(text):
        # 마커 앞 일반 텍스트
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        nums = _expand(m.group(1))
        if nums and all(n in refs_by_num for n in nums):
            instr = _build_encite(nums, refs_by_num, m.group(0))
            _add_field(p, instr, m.group(0))
        else:
            # 인용 ref가 매핑 없으면 텍스트 fallback
            p.add_run(m.group(0))
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


# ── 메인 ──────────────────────────────────────────────────────────────────────

def main():
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "data/exports/ZCB_FINAL.md")
    base = sys.argv[2] if len(sys.argv) > 2 else "ZCB_FINAL_endnote"
    md = src.read_text(encoding="utf-8")

    mt = re.search(r"(?m)^#\s+(.+?)\s*$", md); title = mt.group(1).strip()
    mref = re.search(r"(?im)^##\s*references?\s*$", md)
    body = md[:mref.start()]
    ref_block = md[mref.end():]

    # 36개 ref 파싱 → Reference 객체 → 번호 매핑
    refs_by_num = {}
    ref_lines = []
    for line in ref_block.splitlines():
        s = line.strip()
        mnum = re.match(r"^(\d+)\.\s*(.+)$", s)
        if not mnum:
            continue
        num = int(mnum.group(1))
        entry = mnum.group(2)
        ref = _parse_vancouver_entry(entry)
        refs_by_num[num] = ref
        ref_lines.append((num, format_vancouver(ref, num)))

    # 섹션 파싱
    parts = re.split(r"(?m)^(#{1,3})\s+(.+?)\s*$", body)
    # parts[0] = 첫 # 앞 텍스트, 이후 (hashes, header, content) 트리플
    doc = Document()
    doc.add_heading(title, 0)

    # 첫 # (title)은 이미 처리됨 → 이후 헤더만 순회
    # re.split 결과: ['', '# ', 'Title', 'after title text', '## ', 'Abstract', '...']
    # 이미 # title은 doc.add_heading으로 처리, 나머지 ##/### 헤더는 아래
    i = 1
    while i < len(parts) - 2:
        hashes = parts[i]
        header = parts[i + 1].strip()
        content = parts[i + 2]
        level = len(hashes)
        i += 3
        if level == 1:  # 이미 처리
            continue
        if level == 2:
            doc.add_heading(header, level=1)
        elif level == 3:
            doc.add_heading(header, level=2)
        # content는 일반 문단들 — 빈줄 기준 분리
        for para in re.split(r"\n\s*\n", content):
            p = para.strip()
            if not p:
                continue
            _emit_paragraph(doc, p, refs_by_num)

    # References (필드 없이 plain — EndNote가 cite 필드 기반으로 재포맷 가능)
    doc.add_heading("References", level=1)
    for num, line in sorted(ref_lines):
        doc.add_paragraph(line)

    out = Path(f"data/exports/{base}.docx")
    buf = io.BytesIO(); doc.save(buf); out.write_bytes(buf.getvalue())

    # 검증
    import zipfile
    with zipfile.ZipFile(io.BytesIO(buf.getvalue())) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    n_fld_begin = xml.count('w:fldCharType="begin"')
    n_fld_end = xml.count('w:fldCharType="end"')
    n_encite = xml.count("ADDIN EN.CITE")
    has_record = "&lt;record&gt;" in xml
    print(f"입력: {len(refs_by_num)}개 ref · 본문 ref 매핑: {sorted(refs_by_num)[0]}~{sorted(refs_by_num)[-1]}")
    print(f"출력 docx: {len(out.read_bytes())}B → {out}")
    print(f"필드 begin/end: {n_fld_begin}/{n_fld_end} (일치={n_fld_begin==n_fld_end}) · ADDIN EN.CITE 개수: {n_encite}")
    print(f"travelling library record 임베드: {has_record}")
    if n_fld_begin != n_fld_end or n_encite == 0:
        print("⚠️ 필드 구조 이상")
    else:
        print(f"✅ CWYW 필드 구조 well-formed — EndNote 'Update Citations & Bibliography'로 인식 가능 예상")
        print("   (실제 EndNote 인식 최종 확인은 사용자 측 필요 — 이 환경엔 EndNote 없음)")


if __name__ == "__main__":
    main()
