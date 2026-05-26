"""Citation Workflow E2E — 레퍼런스 풀셋 워크플로를 '실제 실행'으로 검증.

검증 항목:
  ① parse_reference_input  — 붙여넣은 목록 파싱
  ② screen_applicability   — 차용 가능성 임베딩 검수(관련 ref만 usable)
  ③ place_citations        — 본문에 [n] 넘버링 삽입(등장순 번호)
  ④ reference_list_markdown— 번호순 Vancouver 목록
  ⑤ build_cited_docx       — Word 바이트(본문+인용+References)
  ⑥ endnote_bytes/bibtex   — EndNote XML / BibTeX 출력
  ⑦ resolve_references     — 실 PMID PubMed fetch (네트워크 있을 때만)

LLM-무관(임베딩) → Gemini 쿼터와 무관하게 항상 검증 가능.
실행: python scripts/e2e_citation.py
"""
from __future__ import annotations
import io, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

RESULTS = []


def check(name, fn):
    try:
        ok, detail = fn()
        RESULTS.append((name, ok if ok is None else bool(ok), str(detail)[:160]))
    except Exception as e:
        import traceback
        RESULTS.append((name, False, f"EXC {type(e).__name__}: {str(e)[:120]}"))
        traceback.print_exc()


def main():
    import warnings
    warnings.filterwarnings("ignore")
    from src.config.env import bootstrap
    bootstrap()

    from src.export.reference_library import Reference
    from src.export import citation_workflow as cw

    # 내 논문(청소년 ZCB·우울) — 섹션 본문
    paper_sections = {
        "introduction": (
            "Zero-calorie beverage consumption has risen sharply among Korean adolescents "
            "as sugar-tax debates intensify. Prior work links sugar-sweetened beverages to "
            "obesity, but the mental-health correlates of artificial sweeteners remain unclear. "
            "Sleep deprivation is also pervasive in this age group and may confound such associations. "
            "We therefore examined the association between zero-calorie beverage intake and depressive symptoms."
        ),
        "discussion": (
            "Our findings indicate a dose-response association between zero-calorie beverage intake and "
            "depressive symptoms among adolescents. This aligns with reports of gut-brain axis disruption "
            "by non-nutritive sweeteners. Residual confounding by sleep and physical activity cannot be excluded. "
            "Future longitudinal studies should clarify causality."
        ),
    }
    paper_text = " ".join(paper_sections.values())

    # 레퍼런스 후보: 관련 4 + 무관 2 (차용 검수가 무관 ref를 걸러야 함)
    refs = [
        Reference(pmid="1", title="Artificial sweetener beverages and depressive symptoms in adolescents",
                  authors=["Kim J", "Lee S"], journal="J Affect Disord", year="2024",
                  volume="350", pages="100-108",
                  abstract="Zero-calorie beverage intake showed a dose-response association with depression in a national adolescent survey."),
        Reference(pmid="2", title="Non-nutritive sweeteners and the gut-brain axis",
                  authors=["Park H"], journal="Nutrients", year="2023", volume="15", pages="2210",
                  abstract="Artificial sweeteners may disrupt gut microbiota and affect mood via the gut-brain axis."),
        Reference(pmid="3", title="Sleep deprivation and adolescent mental health: KYRBS analysis",
                  authors=["Choi Y"], journal="Sleep Med", year="2022", volume="90", pages="55-62",
                  abstract="Short sleep duration was strongly associated with depressive symptoms among Korean adolescents."),
        Reference(pmid="4", title="Sugar-sweetened beverages and obesity in youth",
                  authors=["Han B"], journal="Obesity", year="2021", volume="29", pages="1200-1210",
                  abstract="Sugar-sweetened beverage consumption predicted weight gain in a youth cohort."),
        Reference(pmid="5", title="Surgical outcomes of total knee arthroplasty in the elderly",
                  authors=["Yoon K"], journal="J Orthop Surg", year="2020", volume="28", pages="33-40",
                  abstract="Total knee arthroplasty showed favorable outcomes in patients over 70 years."),
        Reference(pmid="6", title="Quantum dot solar cell efficiency improvements",
                  authors=["Seo D"], journal="Nat Energy", year="2019", volume="4", pages="900-910",
                  abstract="Quantum dot photovoltaics achieved record power conversion efficiency."),
    ]

    # ① 파싱
    def _t_parse():
        txt = "1. 12345678\n2) 10.1016/j.jad.2024.01.001\n- Artificial sweeteners and mood\n;;\n"
        items = cw.parse_reference_input(txt)
        return len(items) == 3 and "12345678" in items, f"{items}"
    check("① parse_reference_input", _t_parse)

    # ② 차용 검수
    screen_holder = {}
    def _t_screen():
        scr = cw.screen_applicability(refs, paper_text, threshold=0.30)
        screen_holder["scr"] = scr
        usable = [s for s in scr if s["usable"]]
        usable_titles = [s["ref"].title[:30] for s in usable]
        # 관련(ZCB·우울·수면) 최소 2개 usable, 무관(무릎/태양전지) 0개 usable
        irrelevant_usable = any("knee" in s["ref"].title.lower() or "solar" in s["ref"].title.lower()
                                for s in usable)
        return (len(usable) >= 2 and not irrelevant_usable,
                f"usable {len(usable)}/{len(scr)} {usable_titles}, 무관통과={irrelevant_usable}")
    check("② screen_applicability(차용검수)", _t_screen)

    # ③ 본문 넘버링 인용 삽입
    place_holder = {}
    def _t_place():
        scr = screen_holder.get("scr", [])
        usable_refs = [s["ref"] for s in scr if s["usable"]]
        new_sections, ordered = cw.place_citations(paper_sections, usable_refs)
        place_holder["sections"] = new_sections
        place_holder["ordered"] = ordered
        import re as _re
        all_text = " ".join(new_sections.values())
        markers = _re.findall(r"\[\d+(?:, \d+)*\]", all_text)
        # 마커가 본문에 실제 삽입됐고, ordered 개수와 번호 연속성 확인
        nums = sorted({int(n) for m in markers for n in _re.findall(r"\d+", m)})
        seq_ok = nums == list(range(1, len(ordered) + 1)) if ordered else False
        return (len(markers) >= 1 and len(ordered) >= 2 and seq_ok,
                f"마커 {markers}, ordered={len(ordered)}, 번호연속={seq_ok}")
    check("③ place_citations(본문 [n] 삽입)", _t_place)

    # ④ 참고문헌 목록
    def _t_reflist():
        ordered = place_holder.get("ordered", [])
        md = cw.reference_list_markdown(ordered)
        lines = [l for l in md.splitlines() if l.strip()]
        return (len(lines) == len(ordered) and md.startswith("1."),
                f"{len(lines)}줄, 첫줄={md.splitlines()[0][:60] if md else ''!r}")
    check("④ reference_list_markdown", _t_reflist)

    # ⑤ Word 풀셋
    def _t_docx():
        ordered = place_holder.get("ordered", [])
        sections = place_holder.get("sections", {})
        data = cw.build_cited_docx("청소년 무열량음료와 우울", sections, ordered)
        # docx 안에 References 헤딩 + 본문 인용 마커가 실제 들어갔는지 zip 파싱
        ok_struct = False
        try:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                xml = z.read("word/document.xml").decode("utf-8", "replace")
            ok_struct = "References" in xml and "[" in xml
        except Exception:
            pass
        return len(data) > 1500 and ok_struct, f"{len(data)}bytes, References&마커포함={ok_struct}"
    check("⑤ build_cited_docx(Word 풀셋)", _t_docx)

    # ⑥ EndNote / BibTeX
    def _t_export():
        ordered = place_holder.get("ordered", [])
        en = cw.endnote_bytes(ordered)
        bib = cw.bibtex_bytes(ordered)
        en_ok = b"<xml>" in en or b"<records>" in en or b"<record>" in en
        bib_ok = b"@article" in bib.lower() or b"@" in bib
        return (len(en) > 200 and len(bib) > 100 and en_ok and bib_ok,
                f"EndNote={len(en)}b({en_ok}) BibTeX={len(bib)}b({bib_ok})")
    check("⑥ endnote_bytes / bibtex_bytes", _t_export)

    # ⑦ 실 PMID resolve (네트워크 의존)
    def _t_resolve():
        try:
            resolved = cw.resolve_references(["33069327"])  # 실존 PMID
        except Exception as ex:
            return None, f"네트워크 불가(스킵): {ex}"
        if not resolved or not resolved[0].title:
            return None, "네트워크 불가/응답없음(스킵)"
        return True, f"PMID 33069327 → {resolved[0].title[:60]!r} ({resolved[0].journal})"
    check("⑦ resolve_references(실 PubMed)", _t_resolve)

    # ⑧ DOCX 기존 레퍼런스/넘버링 추출
    docx_holder = {}
    def _t_docx_extract():
        from docx import Document
        doc = Document()
        doc.add_heading("Zero-calorie beverages and adolescent depression", 0)
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Zero-calorie beverage intake is rising among adolescents [1]. "
                          "Sleep deprivation may confound this association [2].")
        doc.add_heading("Discussion", level=1)
        doc.add_paragraph("Our dose-response finding aligns with gut-brain axis reports [1,2].")
        doc.add_heading("References", level=1)
        doc.add_paragraph("1. Kim J, Lee S. Artificial sweetener beverages and depression in adolescents. "
                          "J Affect Disord. 2024;350:100-108.")
        doc.add_paragraph("2. Choi Y. Sleep deprivation and adolescent mental health: KYRBS analysis. "
                          "Sleep Med. 2022;90:55-62.")
        buf = io.BytesIO(); doc.save(buf)
        docx_holder["bytes"] = buf.getvalue()
        from src.export import citation_workflow as _cwf
        txt = _cwf.extract_docx_text(docx_holder["bytes"])
        ex = _cwf.extract_existing_references(txt)
        ok = (ex["ref_found"] and len(ex["refs"]) == 2 and ex["markers"] == [1, 2]
              and "Affect" in (ex["refs"][0].journal + ex["refs"][0].title)
              and ex["refs"][0].year == "2024")
        return ok, f"refs={len(ex['refs'])} markers={ex['markers']} y0={ex['refs'][0].year if ex['refs'] else '-'} j0={ex['refs'][0].journal[:20] if ex['refs'] else '-'}"
    check("⑧ DOCX 기존 레퍼런스/넘버링 추출", _t_docx_extract)

    # ⑨ DOCX 풀라운드트립 (LLM-무관: restyle=False) → 첨삭본 docx + EndNote
    def _t_docx_roundtrip():
        from src.export import citation_workflow as _cwf
        r = _cwf.revise_docx_fullset(docx_holder.get("bytes", b""), restyle=False)
        if r.get("error"):
            return False, r["error"]
        en_ok = b"<record>" in r["endnote"] or b"<records>" in r["endnote"]
        # 출력 docx 안에 References 헤딩 + 본문 [n] 마커 보존 확인
        struct = False
        try:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(r["docx"])) as z:
                xml = z.read("word/document.xml").decode("utf-8", "replace")
            struct = "References" in xml and "[1" in xml
        except Exception:
            pass
        return (r["n_refs"] == 2 and len(r["docx"]) > 1500 and en_ok and struct,
                f"n_refs={r['n_refs']} docx={len(r['docx'])}b endnote={len(r['endnote'])}b struct={struct} note={r['note'][:40]}")
    check("⑨ DOCX 풀라운드트립(첨삭본+EndNote)", _t_docx_roundtrip)

    # ── 리포트 ──
    scored = [(n, ok) for n, ok, _ in RESULTS if ok is not None]
    n_ok = sum(1 for _, ok in scored if ok)
    print("\n" + "=" * 66)
    print(f"  Citation Workflow E2E: {n_ok}/{len(scored)} PASS (네트워크 항목 별도)")
    print("=" * 66)
    for name, ok, detail in RESULTS:
        mark = "✅" if ok else ("⚠️ " if ok is None else "❌")
        print(f"  {mark} {name}")
        print(f"       {detail}")
    return 0 if n_ok == len(scored) else 1


if __name__ == "__main__":
    raise SystemExit(main())
