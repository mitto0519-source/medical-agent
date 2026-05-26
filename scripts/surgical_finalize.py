"""수술적 마무리(LLM 미사용, 결정론적):
  1) 딥패스가 환각한 '가짜 Results 블록 #1'(54,170/94.1%/41.3% 등) 제거 → 실제 Results만
  2) 중복 '## Methods' 헤더 1개로
  3) 레퍼런스 37·38 삭제 + 본문 [37]/[38] 마커 제거(범위 보존)
  4) Word/EndNote/BibTeX 재빌드 + 숫자 무결성 감사
실행: python scripts/surgical_finalize.py <in.md> <out_base>
"""
from __future__ import annotations
import io, re, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

DROP = {"37", "38"}


def strip_cites(text: str) -> str:
    def fix(m):
        toks = [t.strip() for t in m.group(1).split(",")]
        keep = [t for t in toks if t not in DROP]
        return ("[" + ", ".join(keep) + "]") if keep else ""
    s = re.sub(r"\[([\d,\s\-]+)\]", fix, text)
    s = re.sub(r"[ \t]+([.,;])", r"\1", s)   # 마커 제거로 생긴 " ." 정리
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s


def main():
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "data/exports/ZCB_final_v2.md")
    base = sys.argv[2] if len(sys.argv) > 2 else "ZCB_FINAL"
    md = src.read_text(encoding="utf-8")

    # 1) 가짜 Results 블록 #1 제거 (## Results 가 2개면 첫 블록 삭제)
    res_idx = [m.start() for m in re.finditer(r"(?im)^##\s*results\s*$", md)]
    if len(res_idx) >= 2:
        md = md[:res_idx[0]] + md[res_idx[1]:]

    # 2) 중복 ## Methods 헤더 축약
    md = re.sub(r"(?im)(^##\s*Methods\s*$\s*){2,}", "## Methods\n\n", md)

    # 3) 본문 [37]/[38] 제거 (References 앞부분만)
    mref = re.search(r"(?im)^##\s*references?\s*$", md)
    body, reftail = md[:mref.start()], md[mref.start():]
    body = strip_cites(body)
    # 참고문헌 37·38 줄 삭제
    ref_lines = []
    for l in reftail.splitlines():
        if re.match(r"^\s*(37|38)\.", l):
            continue
        ref_lines.append(l)
    reftail = "\n".join(ref_lines)
    md = body.rstrip() + "\n\n" + reftail.strip() + "\n"
    Path(f"data/exports/{base}.md").write_text(md, encoding="utf-8")

    # 4) 재빌드
    from src.export.reference_library import to_endnote_xml, to_bibtex
    from src.export.citation_workflow import _parse_vancouver_entry
    mt = re.search(r"(?m)^#\s+(.+?)\s*$", md); title = mt.group(1).strip()
    mref = re.search(r"(?im)^##\s*references?\s*$", md)
    bd = md[:mref.start()]
    reflines = [l.rstrip() for l in md[mref.end():].splitlines() if re.match(r"^\d+\.", l.strip())]
    parts = re.split(r"(?m)^##\s+(.+?)\s*$", bd)
    secs = [[parts[i].strip(), (parts[i + 1] or "").strip()] for i in range(1, len(parts) - 1, 2)]

    from docx import Document
    doc = Document(); doc.add_heading(title, 0)
    lab = {"abstract": "Abstract", "introduction": "Introduction", "methods": "Methods",
           "results": "Results", "discussion": "Discussion"}
    for l, c in secs:
        doc.add_heading(lab.get(l.lower(), l), level=1)
        for para in c.split("\n\n"):
            p = para.strip()
            if not p:
                continue
            if p.startswith("### "):
                doc.add_heading(p[4:].strip(), level=2)
            else:
                doc.add_paragraph(p)
    doc.add_heading("References", level=1)
    for line in reflines:
        doc.add_paragraph(line)
    b = io.BytesIO(); doc.save(b); Path(f"data/exports/{base}.docx").write_bytes(b.getvalue())

    allrefs = [_parse_vancouver_entry(re.sub(r"^\s*\d+[.)]\s*", "", l)) for l in reflines]
    Path(f"data/exports/{base}.xml").write_bytes(to_endnote_xml(allrefs).encode("utf-8"))
    Path(f"data/exports/{base}.bib").write_bytes(to_bibtex(allrefs).encode("utf-8"))

    # 감사
    cites = {int(n) for g in re.findall(r"\[([\d,\s\-]+)\]", bd) for n in re.findall(r"\d+", g)}
    for m in re.findall(r"\[(\d+)\s*-\s*(\d+)", bd):
        cites |= set(range(int(m[0]), int(m[1]) + 1))
    fabs = ["54,170", "94.1", "50.8", "41.3", "11.2", "1.16", "1.79", "0.002", "0.004"]
    leftover = [f for f in fabs if re.search(r"(?<![0-9.])" + re.escape(f), bd)]
    print(f"## Results 수: {len(re.findall(r'(?im)^##\\s*results\\s*$', md))} · ## Methods 수: {len(re.findall(r'(?im)^##\\s*methods\\s*$', md))}")
    print(f"참고문헌 {len(reflines)}개 · 본문 인용 1~{max(cites)} 누락={sorted(set(range(1,len(reflines)+1))-cites)}")
    print(f"가짜숫자 잔존: {leftover if leftover else '없음 ✅'}")
    print(f"저장: {base}.md/.docx/.xml/.bib → data/exports/")


if __name__ == "__main__":
    main()
