"""조유선 첨삭본 + PubMed 신규 레퍼런스 보강 → 본문 [n] 삽입 + 참고문헌 추가 + Word/EndNote 풀셋.

- 기존 28개 참고문헌은 원문 그대로 보존(재포맷 안 함), 신규는 29번부터 이어서 추가.
- 신규 인용은 Introduction/Discussion에서 의미상 가장 가까운 문장 뒤에 [n] 삽입(임베딩).
- EndNote XML/BibTeX는 기존+신규 전체를 Reference 객체로 재구성해 import 호환.
실행: python scripts/augment_and_build.py <restyled.md> <out_basename>
"""
from __future__ import annotations
import io, re, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

QUERIES = [
    "artificial sweetener depression adolescents",
    "non-nutritive sweetener mental health youth",
    "diet beverage depression cohort",
    "sugar-sweetened beverage adolescent depression",
    "ultra-processed food adolescent depression",
    "soft drink depression gut microbiome",
]
THRESHOLD = 0.45   # "의미있는" 차용 기준
MAX_NEW = 10


def main():
    import warnings; warnings.filterwarnings("ignore")
    from src.config.env import bootstrap; bootstrap()
    from src.export.reference_library import (
        search_pubmed, _fetch_pubmed_xml, _parse_pubmed_xml, format_vancouver,
        to_endnote_xml, to_bibtex,
    )
    from src.export.citation_workflow import _parse_vancouver_entry, _emb, _cos

    src = Path(sys.argv[1] if len(sys.argv) > 1 else "data/exports/ZCB_yoosun_claude.md")
    base = sys.argv[2] if len(sys.argv) > 2 else "ZCB_final_yoosun"
    text = src.read_text(encoding="utf-8")

    mt = re.search(r"(?m)^#\s+(.+?)\s*$", text)
    title = mt.group(1).strip()
    mref = re.search(r"(?im)^##\s*references?\s*$", text)
    body = text[:mref.start()]
    ref_lines = [l.rstrip() for l in text[mref.end():].splitlines() if l.strip()]
    n_existing = len(ref_lines)

    # 본문 섹션 파싱
    parts = re.split(r"(?m)^##\s+(.+?)\s*$", body)
    head = parts[0]
    secs = [[parts[i].strip(), (parts[i + 1] or "").rstrip()] for i in range(1, len(parts) - 1, 2)]
    sec_by = {s[0].lower(): idx for idx, s in enumerate(secs)}

    # ── PubMed 신규 후보 수집 + 중복 제거 ──
    existing_lc = " ".join(ref_lines).lower()
    cand, seen = [], set()
    for q in QUERIES:
        try:
            for r in _parse_pubmed_xml(_fetch_pubmed_xml(search_pubmed(q, max_results=4))):
                if not r.pmid or r.pmid in seen:
                    continue
                seen.add(r.pmid)
                if (r.title or "").lower()[:40] in existing_lc:
                    continue
                cand.append(r)
        except Exception as e:
            print(f"  [query 실패] {q}: {str(e)[:70]}")

    # ── 본문(Introduction+Discussion) 문장 후보로 차용 검수/배치 ──
    paper_text = " ".join(s[1] for s in secs)
    pv = _emb(paper_text[:2000])
    scored = []
    for r in cand:
        rt = f"{r.title}. {r.abstract}".strip()[:2000]
        scored.append((_cos(pv, _emb(rt)) if rt.strip(". ") else 0.0, r))
    scored.sort(key=lambda x: -x[0])
    chosen = [(sc, r) for sc, r in scored if sc >= THRESHOLD][:MAX_NEW]
    print(f"PubMed 후보 {len(cand)}개 → 의미있는(≥{THRESHOLD}) 신규 {len(chosen)}개 채택")

    # 인용 배치 대상 문장 = Introduction/Discussion 문장만 (선행연구 인용 위치)
    targets = []  # (sec_idx, sentence, vec)
    for nm in ("introduction", "discussion"):
        if nm in sec_by:
            si = sec_by[nm]
            for sent in re.split(r"(?<=[.])\s+", secs[si][1]):
                if len(sent.strip()) > 40:
                    targets.append((si, sent.strip()))
    tvecs = [_emb(t[1][:400]) for t in targets]

    # 각 신규 ref → 가장 가까운 문장
    placements = []  # (target_index, ref)
    for sc, r in chosen:
        rv = _emb(f"{r.title}. {r.abstract}".strip()[:400])
        bi = max(range(len(targets)), key=lambda j: _cos(rv, tvecs[j]))
        placements.append((bi, r))

    # 등장 순서로 번호 부여 (기존 n_existing 다음부터)
    placements.sort(key=lambda x: x[0])
    new_refs_ordered = []
    by_target = {}
    for bi, r in placements:
        num = n_existing + len(new_refs_ordered) + 1
        new_refs_ordered.append(r)
        by_target.setdefault(bi, []).append(num)

    # 문장 뒤에 마커 삽입
    for bi, nums in by_target.items():
        si, sent = targets[bi]
        marker = "[" + ", ".join(str(n) for n in sorted(nums)) + "]"
        secs[si][1] = secs[si][1].replace(sent, f"{sent.rstrip()} {marker}", 1)

    # 참고문헌 줄: 기존 그대로 + 신규(format_vancouver, 번호 연속)
    for i, r in enumerate(new_refs_ordered, start=n_existing + 1):
        ref_lines.append(format_vancouver(r, i))

    # ── 최종 md 저장 ──
    new_md = head.rstrip() + "\n\n" + "\n\n".join(f"## {l}\n{c}" for l, c in secs) + \
        "\n\n## References\n" + "\n".join(ref_lines) + "\n"
    Path(f"data/exports/{base}.md").write_text(new_md, encoding="utf-8")

    # ── DOCX (기존 ref 원문 보존 + 신규 추가) ──
    from docx import Document
    doc = Document()
    doc.add_heading(title, 0)
    label_map = {"abstract": "Abstract", "introduction": "Introduction",
                 "methods": "Methods", "results": "Results", "discussion": "Discussion"}
    for l, c in secs:
        doc.add_heading(label_map.get(l.lower(), l), level=1)
        for para in c.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    doc.add_heading("References", level=1)
    for line in ref_lines:
        doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf)
    Path(f"data/exports/{base}.docx").write_bytes(buf.getvalue())

    # ── EndNote XML / BibTeX (기존 28 파싱 + 신규 = 전체) ──
    all_refs = [_parse_vancouver_entry(re.sub(r"^\s*\d+[.)]\s*", "", l)) for l in ref_lines]
    Path(f"data/exports/{base}.xml").write_bytes(to_endnote_xml(all_refs).encode("utf-8"))
    Path(f"data/exports/{base}.bib").write_bytes(to_bibtex(all_refs).encode("utf-8"))

    print(f"\n신규 인용 배치 ({len(new_refs_ordered)}개):")
    for i, r in enumerate(new_refs_ordered, start=n_existing + 1):
        au = r.authors[0] if r.authors else "?"
        print(f"  [{i}] {au} et al. {r.title[:80]} ({r.journal} {r.year})")
    print(f"\n참고문헌 총 {len(ref_lines)}개 (기존 {n_existing} + 신규 {len(new_refs_ordered)})")
    print(f"출력: {base}.md / .docx / .xml / .bib → data/exports/")


if __name__ == "__main__":
    main()
