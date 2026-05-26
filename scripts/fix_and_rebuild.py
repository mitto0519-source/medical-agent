"""인접 인용 마커 병합([10][29]→[10, 29]) + Word/EndNote/BibTeX 재빌드 (참고문헌 원문 보존).

LLM 미사용 — 배치는 그대로 두고 표기/포맷만 정리. 실행: python scripts/fix_and_rebuild.py <md> <base>
"""
from __future__ import annotations
import io, re, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")


def merge_brackets(s: str) -> str:
    # 인접 대괄호 인용 병합: "[10][29]" / "[15-17][33, 34]" → "[10, 29]" / "[15-17, 33, 34]"
    prev = None
    while prev != s:
        prev = s
        s = re.sub(r"\]\s*\[", ", ", s)
    # 문장부호 뒤 인용 → 부호 앞으로 (문서 표준: "...effect [29, 35]."): effect.[29,35]→effect [29, 35].
    s = re.sub(r"(\w)\s*([.!?])\s*(\[[\d,\s\-]+\])", r"\1 \3\2", s)
    s = re.sub(r"([A-Za-z\*])(\[\d)", r"\1 \2", s)  # 단어 뒤 공백 없는 인용: adults[35]→adults [35]
    return s


def main():
    from src.export.reference_library import to_endnote_xml, to_bibtex
    from src.export.citation_workflow import _parse_vancouver_entry

    src = Path(sys.argv[1] if len(sys.argv) > 1 else "data/exports/ZCB_final_yoosun.md")
    base = sys.argv[2] if len(sys.argv) > 2 else "ZCB_final_yoosun"
    text = src.read_text(encoding="utf-8")

    mt = re.search(r"(?m)^#\s+(.+?)\s*$", text); title = mt.group(1).strip()
    mref = re.search(r"(?im)^##\s*references?\s*$", text)
    body = text[:mref.start()]
    ref_lines = [l.rstrip() for l in text[mref.end():].splitlines() if l.strip()]

    parts = re.split(r"(?m)^##\s+(.+?)\s*$", body)
    head = parts[0].rstrip()
    secs = [[parts[i].strip(), merge_brackets((parts[i + 1] or "").rstrip())]
            for i in range(1, len(parts) - 1, 2)]

    new_md = head + "\n\n" + "\n\n".join(f"## {l}\n{c}" for l, c in secs) + \
        "\n\n## References\n" + "\n".join(ref_lines) + "\n"
    Path(f"data/exports/{base}.md").write_text(new_md, encoding="utf-8")

    from docx import Document
    doc = Document(); doc.add_heading(title, 0)
    lab = {"abstract": "Abstract", "introduction": "Introduction", "methods": "Methods",
           "results": "Results", "discussion": "Discussion"}
    for l, c in secs:
        doc.add_heading(lab.get(l.lower(), l), level=1)
        for para in c.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    doc.add_heading("References", level=1)
    for line in ref_lines:
        doc.add_paragraph(line)
    b = io.BytesIO(); doc.save(b); Path(f"data/exports/{base}.docx").write_bytes(b.getvalue())

    all_refs = [_parse_vancouver_entry(re.sub(r"^\s*\d+[.)]\s*", "", l)) for l in ref_lines]
    Path(f"data/exports/{base}.xml").write_bytes(to_endnote_xml(all_refs).encode("utf-8"))
    Path(f"data/exports/{base}.bib").write_bytes(to_bibtex(all_refs).encode("utf-8"))

    body_all = "\n".join(c for _, c in secs)
    marks = sorted({int(n) for g in re.findall(r"\[([\d,\s\-]+)\]", body_all) for n in re.findall(r"\d+", g)})
    print(f"참고문헌 {len(ref_lines)}개 · 본문 인용 마커 {marks[:6]}...{marks[-4:]} (최대 {max(marks)})")
    print(f"재빌드: {base}.md/.docx/.xml/.bib → data/exports/")


if __name__ == "__main__":
    main()
