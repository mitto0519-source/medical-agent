"""Paper Writer — 조유선 스타일로 논문 초안 생성

AuthorProfile(스타일 시드) + MethodsLibrary + DatasetLibrary + RAG컨텍스트를
결합해 완성도 높은 논문 초안을 생성.

Upgrade (2026-05): real stat injection from StatBridge + DOCX export + journal templates.
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional

from src.config.logging_config import get_logger
from src.llm import get_llm_client

_log = get_logger(__name__)


def format_vancouver(paper: dict, index: int = 1) -> str:
    """Format a paper dict as a Vancouver-style reference.

    paper keys: title, authors (list|str), journal, year, volume, issue, pages,
                pmid, doi
    """
    authors = paper.get("authors") or []
    if isinstance(authors, list):
        if len(authors) > 6:
            author_str = ", ".join(authors[:6]) + ", et al"
        else:
            author_str = ", ".join(authors)
    else:
        author_str = str(authors)

    title = paper.get("title", "").rstrip(".")
    journal = paper.get("journal", "") or paper.get("source", "")
    year = paper.get("year", "") or paper.get("pub_date", "")
    volume = paper.get("volume", "")
    issue = paper.get("issue", "")
    pages = paper.get("pages", "") or paper.get("page", "")
    doi = paper.get("doi", "")
    pmid = paper.get("pmid", "")

    vol_issue = f"{volume}" + (f"({issue})" if issue else "")
    location = f"{vol_issue}:{pages}".strip(":") if (vol_issue or pages) else ""

    parts = []
    if author_str:
        parts.append(author_str + ".")
    if title:
        parts.append(title + ".")
    if journal:
        journal_part = f"{journal}." + (f" {year};" if year else "")
        if location:
            journal_part += location + "."
        parts.append(journal_part)
    elif year:
        parts.append(f"{year}.")
    if doi:
        parts.append(f"doi:{doi}")
    elif pmid:
        parts.append(f"PMID:{pmid}")

    ref = f"{index}. " + " ".join(parts)
    return re.sub(r"\s+", " ", ref).strip()


def format_references_vancouver(papers: list) -> str:
    """Return a numbered Vancouver reference list from a list of paper dicts."""
    return "\n".join(format_vancouver(p, i + 1) for i, p in enumerate(papers))


class PaperWriter:
    """저자 스타일 시드를 적용한 논문 작성기.

    사용 흐름:
        writer = PaperWriter(author_profile, methods_lib, dataset_lib)
        draft = writer.write_abstract(topic, study_info, results)
        full  = writer.write_full_paper(topic, study_info, results, references)
    """

    def __init__(
        self,
        author_profile,          # AuthorProfile instance
        methods_library=None,    # MethodsLibrary instance
        dataset_library=None,    # DatasetLibrary instance
        rag_pipeline=None,       # RAGPipeline for reference retrieval
        llm_client=None,
        api_key: Optional[str] = None,
    ):
        self._profile = author_profile
        self._methods = methods_library
        self._datasets = dataset_library
        # ★ FIX 10 (회로 끊김): rag_pipeline=None이면 자동으로 로컬 ChromaDB(20,894 chunks)
        #   를 붙임. 호출자가 인자 전달 안 해도 자산이 회로에 연결되도록.
        if rag_pipeline is None:
            try:
                from src.rag.pipeline import RAGPipeline
                rag_pipeline = RAGPipeline()
                n = 0
                try:
                    n = rag_pipeline._store.count() if hasattr(rag_pipeline._store, "count") else 0
                except Exception:
                    pass
                _log.info("[PaperWriter] auto-attached RAGPipeline (chunks≈%s)", n)
            except Exception as _e:
                _log.warning("[PaperWriter] RAG auto-attach 실패 → reference 없이 진행: %s", _e)
                rag_pipeline = None
        self._rag = rag_pipeline
        self._client = llm_client or get_llm_client(api_key=api_key)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def write_abstract(
        self,
        topic: str,
        background: str,
        objective: str,
        methods_summary: str,
        results_summary: str,
        conclusion: str,
        word_limit: int = 250,
    ) -> str:
        """저자 스타일로 Abstract 작성."""

        system = self._profile.get_system_prompt()
        prompt = f"""Write a structured abstract for the following medical research paper.
Word limit: {word_limit} words.
Use this author's exact style and formatting.

TOPIC: {topic}

SECTIONS TO INCLUDE:
Background: {background}
Objective: {objective}
Methods: {methods_summary}
Results: {results_summary}
Conclusion: {conclusion}

Write the complete abstract now. Use the standard structured format (Background / Objective / Methods / Results / Conclusion).
"""
        return self._generate(system, prompt)

    def write_introduction(
        self,
        topic: str,
        background_facts: List[str],
        knowledge_gap: str,
        hypothesis: str,
        reference_context: Optional[str] = None,
    ) -> str:
        """저자 스타일로 Introduction 작성."""

        ref_block = f"\n\nREFERENCE CONTEXT FROM LITERATURE:\n{reference_context}" if reference_context else ""
        system = self._profile.get_system_prompt()
        prompt = f"""Write the Introduction section for this medical research paper.
Match the author's exact introduction structure and style.

TOPIC: {topic}

KEY BACKGROUND FACTS:
{chr(10).join(f'- {f}' for f in background_facts)}

KNOWLEDGE GAP: {knowledge_gap}

STUDY HYPOTHESIS/AIM: {hypothesis}
{ref_block}

Write the full Introduction section. Build from broad context → specific gap → study aim.
"""
        return self._generate(system, prompt)

    def write_methods(
        self,
        study_design: str,
        population: str,
        dataset_name: Optional[str] = None,
        exposure: str = "",
        outcome: str = "",
        covariates: str = "",
        statistical_methods: List[str] = None,
    ) -> str:
        """저자 스타일로 Methods 섹션 작성."""

        dataset_ctx = ""
        if dataset_name and self._datasets:
            dataset_ctx = self._datasets.get_context(dataset_name)

        methods_ctx = ""
        if statistical_methods and self._methods:
            methods_ctx = self._methods.get_context_for_claude(statistical_methods)

        system = self._profile.get_system_prompt()
        prompt = f"""Write the Methods section for this medical research paper.
Use this author's level of detail and formatting exactly.

STUDY DESIGN: {study_design}
POPULATION: {population}
EXPOSURE: {exposure}
OUTCOME: {outcome}
COVARIATES: {covariates}

{dataset_ctx}

{methods_ctx}

Write the full Methods section including: Study Design and Population, Exposure Assessment,
Outcome Assessment, Covariate Assessment, Statistical Analysis.
"""
        return self._generate(system, prompt)

    def write_results(
        self,
        descriptive_stats: str,
        main_findings: List[str],
        subgroup_findings: Optional[str] = None,
        sensitivity_findings: Optional[str] = None,
    ) -> str:
        """저자 스타일로 Results 섹션 작성."""

        sub_block = f"\nSubgroup findings:\n{subgroup_findings}" if subgroup_findings else ""
        sens_block = f"\nSensitivity analysis findings:\n{sensitivity_findings}" if sensitivity_findings else ""

        system = self._profile.get_system_prompt()
        prompt = f"""Write the Results section for this medical research paper.
Match the author's exact results-reporting style (how they present ORs, HRs, CIs, p-values, etc.).

DESCRIPTIVE STATISTICS:
{descriptive_stats}

MAIN FINDINGS:
{chr(10).join(f'- {f}' for f in main_findings)}
{sub_block}
{sens_block}

Write the full Results section. Describe participant characteristics first, then main analyses,
then subgroup/sensitivity analyses. Use exact numeric reporting format this author uses.
"""
        return self._generate(system, prompt)

    def write_discussion(
        self,
        main_findings: List[str],
        comparison_with_literature: str,
        mechanisms: str,
        strengths: List[str],
        limitations: List[str],
        conclusion_message: str,
        reference_context: Optional[str] = None,
    ) -> str:
        """저자 스타일로 Discussion 작성."""

        ref_block = f"\nREFERENCE CONTEXT:\n{reference_context}" if reference_context else ""
        system = self._profile.get_system_prompt()
        prompt = f"""Write the Discussion section for this medical research paper.
Match the author's discussion structure and hedging language exactly.

MAIN FINDINGS TO DISCUSS:
{chr(10).join(f'- {f}' for f in main_findings)}

COMPARISON WITH EXISTING LITERATURE:
{comparison_with_literature}

PROPOSED MECHANISMS:
{mechanisms}

STRENGTHS:
{chr(10).join(f'- {s}' for s in strengths)}

LIMITATIONS:
{chr(10).join(f'- {l}' for l in limitations)}

KEY CONCLUSION MESSAGE: {conclusion_message}
{ref_block}

Write the full Discussion. Structure: summary of findings → comparison with literature →
mechanisms → strengths/limitations → conclusion.
"""
        return self._generate(system, prompt)

    def write_section(
        self,
        section: str,
        topic: str,
        study_info: Dict,
        results: Dict,
    ) -> str:
        """Streamlit UI 섹션별 작성용 통합 어댑터.

        section: 'Abstract' | 'Introduction' | 'Methods' | 'Results' | 'Discussion'
        study_info: dataset, design, sample_size, exposure, outcome, population 등
        results: summary(필수)
        """
        results_text = results.get("summary", "") if isinstance(results, dict) else str(results)
        dataset = study_info.get("dataset", "KYRBS")
        design = study_info.get("design", "Cross-sectional")
        population = study_info.get("population", "Korean adolescents")
        exposure = study_info.get("exposure", "")
        outcome = study_info.get("outcome", "")
        sample_size = study_info.get("sample_size", "")
        journal = study_info.get("journal", "")

        sys_p = self._profile.get_system_prompt()

        if section == "Abstract":
            return self.write_abstract(
                topic=topic,
                background=f"Korean public health study using {dataset}",
                objective=f"To examine the association in {population}",
                methods_summary=f"{design}, {dataset}, n={sample_size}",
                results_summary=results_text,
                conclusion="Findings support public health interventions.",
            )
        elif section == "Introduction":
            return self.write_introduction(
                topic=topic,
                background_facts=[
                    f"Study population: {population}",
                    f"Dataset: {dataset}",
                    f"Exposure: {exposure}",
                    f"Outcome: {outcome}",
                ],
                knowledge_gap="Gap in Korean adolescent population regarding this association",
                hypothesis=f"Exposure is associated with outcome in {population}",
            )
        elif section == "Methods":
            dataset_key = None
            if self._datasets:
                for ds in self._datasets.list_datasets():
                    if ds.upper() in dataset.upper():
                        dataset_key = ds
                        break
            dataset_ctx = self._datasets.get_context(dataset_key) if (dataset_key and self._datasets) else ""
            methods_ctx = self._methods.get_context_for_claude(["logistic_regression", "complex_sampling"]) if self._methods else ""
            return self._generate(sys_p, f"""Write the Methods section.
STUDY DESIGN: {design}
POPULATION: {population}
DATASET: {dataset} (n={sample_size})
EXPOSURE: {exposure}
OUTCOME: {outcome}
COVARIATES: age, sex, socioeconomic status
{dataset_ctx}
{methods_ctx}
Write the full Methods section: Study Design and Population, Variables, Statistical Analysis.""")
        elif section == "Results":
            return self.write_results(
                descriptive_stats=f"n={sample_size}, {dataset}, {design}",
                main_findings=[results_text] if results_text else ["Not provided"],
            )
        elif section == "Discussion":
            return self.write_discussion(
                main_findings=[results_text] if results_text else ["Not provided"],
                comparison_with_literature="Consistent with existing Korean public health literature",
                mechanisms="Proposed biological and behavioral mechanisms",
                strengths=[f"Large nationally representative sample ({dataset})", "Standardized methodology"],
                limitations=["Cross-sectional design limits causal inference", "Self-reported data"],
                conclusion_message="These findings support targeted public health interventions.",
            )
        else:
            raise ValueError(f"Unknown section: {section}. Use Abstract/Introduction/Methods/Results/Discussion")

    def refine_section(
        self,
        section: str,
        current_text: str,
        instruction: str,
        study_info: Optional[Dict] = None,
    ) -> str:
        """기존 섹션 텍스트를 사용자 지시에 따라 개선 — '바이브 논문' 핵심.

        AI가 새로 쓰는 게 아니라, 사람이 쓴(또는 기존) 텍스트를 instruction대로 다듬어
        사람의 작업을 거든다. 통계·데이터·사실은 보존.
        """
        if not current_text or not current_text.strip():
            raise ValueError("개선할 텍스트가 비어 있습니다. 먼저 내용을 입력하거나 AI 작성을 사용하세요.")
        if not instruction or not instruction.strip():
            instruction = "Improve clarity, academic tone, and flow."
        sys_p = self._profile.get_system_prompt()
        ctx = ""
        if study_info:
            ctx = (
                f"\nSTUDY CONTEXT: dataset={study_info.get('dataset', '')}, "
                f"exposure={study_info.get('exposure', '')}, outcome={study_info.get('outcome', '')}, "
                f"population={study_info.get('population', '')}"
            )
        prompt = f"""Improve the following {section} section of a medical research paper.

USER REQUEST: {instruction}{ctx}

CURRENT TEXT:
{current_text}

Apply the user's request while preserving all data, statistics, and factual claims.
Keep the author's academic writing style. Output ONLY the improved section text — no headers, no commentary."""
        return self._generate(sys_p, prompt)

    def write_full_paper(
        self,
        topic: str,
        study_info: Dict,
        results: Dict,
        reference_context: Optional[str] = None,
        ref_lib=None,
        feedback_context: Optional[str] = None,
    ) -> str:
        """전체 논문 초안 한번에 생성.

        study_info: dataset, design, sample_size, survey_year, journal,
                    population, exposure, outcome, covariates, methods_list 등
        results: summary(필수), main_findings(선택), descriptive, subgroup,
                 sensitivity, conclusion, literature_comparison, mechanisms
        """
        # ── results 정규화: summary → main_findings 자동 변환 ──────────
        results_summary = results.get("summary", "")
        main_findings = results.get("main_findings") or (
            [results_summary] if results_summary else ["Not provided"]
        )

        dataset_name = study_info.get("dataset", "")
        design = study_info.get("design", "cross-sectional")
        population = study_info.get("population", "")
        exposure = study_info.get("exposure", "")
        outcome = study_info.get("outcome", "")
        sample_size = study_info.get("sample_size", "")
        journal = study_info.get("journal", "")

        # dataset_name이 긴 서술형이면 "KYRBS" 같은 약어만 추출해 메서드에 넘김
        dataset_key = None
        if self._datasets:
            for ds_name in self._datasets.list_datasets():
                if ds_name.upper() in dataset_name.upper():
                    dataset_key = ds_name
                    break

        sections = {}
        ref_block = f"\n\nREFERENCE CONTEXT FROM LITERATURE:\n{reference_context}" if reference_context else ""
        feedback_block = f"\n\nPAST REVIEWER LESSONS (apply to avoid known pitfalls):\n{feedback_context}" if feedback_context else ""
        sys_p = self._profile.get_system_prompt()
        dataset_ctx = self._datasets.get_context(dataset_key) if (dataset_key and self._datasets) else ""
        methods_ctx = self._methods.get_context_for_claude(
            study_info.get("methods_list", ["logistic_regression"])
        ) if self._methods else ""
        # ★ FIX 2 (라인 430 원본): 'depression'은 outcome인데 covariate에 잘못 들어가 있었음.
        #   STATA 실제 Model 2 covariates 12개로 교체. KYRBS ZCB-Depression 분석 사양.
        covariates = study_info.get(
            "covariates",
            "sex, age category (12-13/14-15/16-18), school level (middle/high), "
            "academic performance (tertile), household SES (tertile), "
            "BMI category (under/normal/overweight-obese), ever smoking, ever drinking, "
            "sugar-sweetened beverage intake (tertile), caffeine intake (tertile), "
            "physical activity (low/moderate/high), breakfast skipping"
        )

        # ── raw_examples from yoosun_cho.json — user_prompt에 직접 박아 few-shot ──
        # system_prompt에만 박으면 LLM이 "참고만 하고 자기 양식으로 fallback"하는 사고 차단.
        # ★ FIX 9 (토큰 절약): 한 번 빌드해 5섹션 재사용. 매 호출마다 빌드/전송 안 함.
        exemplars_block = self._build_exemplars_block()
        # 섹션별 max_tokens 차등 — Abstract/Intro/Results 낮춤, Methods/Discussion 만 길게
        MT = {"abstract": 1800, "introduction": 3500, "methods": 5500,
              "results": 4000, "discussion": 5500}

        # ── ★ FIX 12 (백엔드 자산화 — 유형 카탈로그 mix & match): ──────────
        # intent_sensor가 픽업한 PaperOrientation(novelty/consistency/...)을 읽어
        # 섹션별 typology 패턴 예문을 user_prompt에 박는다. 사용자에게는 노출 X.
        active_orientations: List[str] = []
        try:
            from src.agent.intent_sensor import get_current as _intent_now
            sig = _intent_now()
            if sig:
                # implicit_emphasis에 박힌 orientation imprint label로부터 역추출
                imp = " ".join(sig.implicit_emphasis or [])
                for o in ["novelty", "consistency", "innovation",
                           "public_health", "methodological_rigor"]:
                    if o in (sig.explicit_request or "") or any(
                            kw in imp for kw in {
                                "novelty": ["novel_finding", "first_adolescent_study", "largest_N"],
                                "consistency": ["consistent_across_subgroups", "robust_to_sensitivity"],
                                "innovation": ["novel_mechanism", "new_methodology"],
                                "public_health": ["policy_implication", "clinical_actionable"],
                                "methodological_rigor": ["complex_survey_design", "multi_covariate_adjustment"],
                            }[o]):
                        active_orientations.append(o)
        except Exception:
            pass
        try:
            from src.knowledge.paper_typology import get_typology_block_for_orientation
            typology_blocks: Dict[str, str] = {}
            for sec_key in ("introduction", "methods", "results", "discussion"):
                typology_blocks[sec_key] = get_typology_block_for_orientation(
                    sec_key, active_orientations, n_per_type=2)
            _log.info("[PaperWriter] typology patterns: orientations=%s, picked={%s}",
                       active_orientations,
                       ", ".join(f"{k}:{'Y' if v else 'n'}" for k, v in typology_blocks.items()))
        except Exception as _e:
            _log.warning("[PaperWriter] typology block fail: %s", _e)
            typology_blocks = {}

        # ── ★ FIX 11 (시드 12,301편 회로 연결): 섹션별 query로 RAG retrieval ──
        # data/chromadb의 papers 컬렉션(20,894 chunks)에서 섹션에 맞는 실 발췌를
        # 끌어와 user_prompt에 박는다. 지금까지는 self._rag=None이라 한 줄도 안 박혔음.
        SECTION_QUERIES = {
            "introduction": f"{exposure} {outcome} {population} epidemiology adolescent public health background",
            "methods":      f"{exposure} {outcome} {design} survey-weighted logistic regression covariate adjustment",
            "results":      f"{exposure} {outcome} odds ratio confidence interval subgroup sex stratified",
            "discussion":   f"{exposure} {outcome} mechanism limitation public health policy implication",
        }
        section_rag_blocks: Dict[str, str] = {}
        if self._rag is not None:
            for sec_key, q in SECTION_QUERIES.items():
                try:
                    hits = self._rag.search_multistage(q, n_final=4, n_pool=20) \
                            if hasattr(self._rag, "search_multistage") \
                            else self._rag.search(q, n_results=4)
                except Exception as _e:
                    _log.warning("[PaperWriter] RAG %s search 실패: %s", sec_key, _e)
                    hits = []
                if not hits:
                    section_rag_blocks[sec_key] = ""
                    continue
                # 발췌 N개를 짧게 — 한 발췌 800자, 총 ~3200자 추가 (토큰 ~800)
                pieces = []
                for i, h in enumerate(hits[:4], 1):
                    txt = (h.get("text") or "")[:800].replace("\n", " ").strip()
                    md = h.get("metadata") or {}
                    src_label = md.get("pmcid") or md.get("pmid") or md.get("source") or f"chunk{i}"
                    pieces.append(f"[{src_label}] {txt}")
                section_rag_blocks[sec_key] = (
                    f"\n\n## RELEVANT LITERATURE EXCERPTS (from {len(hits)} retrieved chunks)\n"
                    "Use these passages to ground your claims. Paraphrase, do NOT plagiarize. "
                    "Cite the bracketed source when you draw from a passage.\n\n"
                    + "\n\n---\n\n".join(pieces)
                    + "\n\n"
                )
            _log.info("[PaperWriter] RAG hits per section: %s",
                       {k: ('hit' if v else 'miss') for k, v in section_rag_blocks.items()})

        # ── Step 1: Introduction (참고문헌 컨텍스트 활용, 독립 생성) ─────────
        _log.info("[PaperWriter] Introduction 작성 중...")
        sections["introduction"] = self._generate(sys_p, f"""{exemplars_block}

NOW WRITE the Introduction section for this medical research paper in the EXACT SAME cadence/voice/sentence-rhythm/hedging as the exemplars above. Do not generic-academic.

TOPIC: {topic}
EXPOSURE: {exposure}
OUTCOME: {outcome}
POPULATION: {population}
STUDY DESIGN: {design}{ref_block}{feedback_block}{typology_blocks.get("introduction","")}{section_rag_blocks.get("introduction","")}

Structure: broad public health context → specific problem → knowledge gap → study aim.
Keep strictly to this topic. Do NOT refer to unrelated studies. Write 4–5 paragraphs.""",
            max_tokens=MT["introduction"])

        # ── Step 2: Methods ─────────────────────
        _log.info("[PaperWriter] Methods 작성 중...")
        sections["methods"] = self._generate(sys_p, f"""{exemplars_block}

NOW WRITE the Methods section in the EXACT SAME author voice as the exemplars above (specific verb choice, sentence rhythm, level of methodological detail). Do NOT default to generic academic tone.

TOPIC: {topic}
STUDY DESIGN: {design}
DATASET: {dataset_name} (n={sample_size})
EXPOSURE: {exposure}
OUTCOME: {outcome}
COVARIATES: {covariates}

{dataset_ctx}
{methods_ctx}{typology_blocks.get("methods","")}{section_rag_blocks.get("methods","")}

Write subsections: Study Design and Population / Exposure / Outcome / Covariates / Statistical Analysis.
Include complex survey analysis (stratification, clustering, weights) if KYRBS data.""",
            max_tokens=MT["methods"])

        # ── Step 3: Results ──────────
        _log.info("[PaperWriter] Results 작성 중...")
        methods_snippet = sections["methods"][:600]
        sections["results"] = self._generate(sys_p, f"""{exemplars_block}

NOW WRITE the Results section in the EXACT SAME numeric reporting cadence as exemplars above. Match the author's verb pattern (e.g. "was associated with", "showed", "did not differ"), sentence length variation, and order of presentation.

TOPIC: {topic}
EXPOSURE: {exposure}
OUTCOME: {outcome}
DATASET: {dataset_name} (n={sample_size})

METHODS USED (for consistency):
{methods_snippet}

KEY FINDINGS (use these EXACT numbers — do NOT invent or round):
{chr(10).join(f'- {f}' for f in main_findings)}

Additional subgroup findings: {results.get("subgroup", "Not provided")}
Sensitivity analyses: {results.get("sensitivity", "Not provided")}{typology_blocks.get("results","")}{section_rag_blocks.get("results","")}

Subsections: Participant characteristics → Main analysis → Subgroup → Sensitivity.""",
            max_tokens=MT["results"])

        # ── Step 4: Discussion ──
        _log.info("[PaperWriter] Discussion 작성 중...")
        results_snippet = sections["results"][:800]
        sections["discussion"] = self._generate(sys_p, f"""{exemplars_block}

NOW WRITE the Discussion in the EXACT SAME hedging rhythm and topic-sentence cadence as exemplars. The author opens with the headline finding first, then comparison, then mechanisms, then limitations, then policy implication. Mirror that rhythm.

TOPIC: {topic}
EXPOSURE: {exposure}
OUTCOME: {outcome}

ACTUAL RESULTS SECTION (refer to these exact findings):
{results_snippet}

COMPARISON WITH LITERATURE: {results.get("literature_comparison", "Discuss in relation to prior studies on this topic.")}
MECHANISMS: {results.get("mechanisms", "Propose biologically plausible mechanisms.")}
STRENGTHS: nationwide representative sample, large n={sample_size}, validated survey instrument
LIMITATIONS: cross-sectional design (cannot establish causality), self-reported data, residual confounding
{ref_block}{feedback_block}{typology_blocks.get("discussion","")}{section_rag_blocks.get("discussion","")}

Structure: 1) Summary of main findings  2) Comparison with existing literature
3) Proposed mechanisms  4) Strengths and limitations  5) Public health conclusion""",
            max_tokens=MT["discussion"])

        # ── Step 5: Abstract ─────────
        _log.info("[PaperWriter] Abstract 작성 중 (전 섹션 통합)...")
        sections["abstract"] = self._generate(sys_p, f"""{exemplars_block}

NOW WRITE the structured abstract in EXACT SAME author voice. 300 words max. Structure: Background / Objective / Methods / Results / Conclusion. Match the exemplars' density of numbers and the way Background is opened (typically one specific framing sentence, not a generic 'X is a major public health concern').

TOPIC: {topic}
TARGET JOURNAL: {journal}

METHODS SUMMARY:
{sections["methods"][:500]}

RESULTS SUMMARY:
{sections["results"][:500]}

CONCLUSION FROM DISCUSSION:
{sections["discussion"][-400:]}

Base the abstract STRICTLY on the sections above. Do NOT invent numbers.""",
            max_tokens=MT["abstract"])

        # ★ FIX 15 (자산화 3단계 — Yoosun 최종 변환):
        # 본문 생성 + anti-AI 필터(_generate 내부) 통과 후, 별도 LLM 호출로
        # Yoosun voice로 최종 변환. yoosun_voice=True (study_info에서 또는 기본).
        # 비용 절감 위해 Discussion + Introduction만 변환 (가장 voice가 드러나는 섹션).
        yoosun_finalize_on = bool(study_info.get("yoosun_finalize", True))
        if yoosun_finalize_on:
            try:
                from src.research.yoosun_finalize import finalize as _yfin
                for sec_key, label in [("introduction", "Introduction"),
                                          ("discussion", "Discussion")]:
                    body = sections.get(sec_key) or ""
                    if len(body) < 400:
                        continue
                    _log.info("[PaperWriter] Yoosun finalize: %s (%d자)", label, len(body))
                    new_body = _yfin(body, section_label=label,
                                       llm_client=self._client,
                                       max_tokens=min(len(body) * 2, 6000))
                    if new_body and len(new_body) > len(body) * 0.5:
                        sections[sec_key] = new_body
            except Exception as _e:
                _log.warning("[PaperWriter] Yoosun finalize 실패: %s", _e)

        # 섹션 dict를 인스턴스에 보관 (JournalDocxExporter가 접근할 수 있도록)
        self.last_sections = {
            "Abstract": sections["abstract"],
            "Introduction": sections["introduction"],
            "Methods": sections["methods"],
            "Results": sections["results"],
            "Discussion": sections["discussion"],
        }

        # Assemble (표준 논문 순서: Abstract → Introduction → Methods → Results → Discussion)
        author = self._profile.author_name
        separator = "=" * 70
        paper = f"""{separator}
{topic.upper()}
{separator}
By {author}

{separator}
ABSTRACT
{separator}
{sections['abstract']}

{separator}
INTRODUCTION
{separator}
{sections['introduction']}

{separator}
METHODS
{separator}
{sections['methods']}

{separator}
RESULTS
{separator}
{sections['results']}

{separator}
DISCUSSION
{separator}
{sections['discussion']}
"""
        # ── References 섹션 추가 (ref_lib이 있을 때만) ────────────────────
        if ref_lib is not None:
            try:
                refs = ref_lib.get_refs() if hasattr(ref_lib, "get_refs") else []
                if refs:
                    from src.export.reference_library import format_reference
                    ref_lines = [format_reference(r, "Vancouver", i + 1) for i, r in enumerate(refs)]
                    paper += f"\n{separator}\nREFERENCES\n{separator}\n"
                    paper += "\n".join(ref_lines) + "\n"
                    self.last_sections["References"] = "\n".join(ref_lines)
            except Exception as e:
                _log.warning("References 섹션 추가 실패: %s", e)

        return paper

    # ------------------------------------------------------------------
    # Stat-injected paper generation
    # ------------------------------------------------------------------

    def write_full_paper_with_stats(
        self,
        topic: str,
        study_info: Dict,
        stat_result: dict,
        reference_context: Optional[str] = None,
        ref_lib=None,
        review_result: dict | None = None,
        feedback_context: Optional[str] = None,
    ) -> str:
        """실제 통계 결과(StatBridge)를 주입해 논문 초안 생성.

        stat_result: StatBridge.AnalysisResult.to_dict()
        """
        results = self._stat_to_results(stat_result)
        paper = self.write_full_paper(
            topic, study_info, results,
            reference_context=reference_context,
            ref_lib=ref_lib,
            feedback_context=feedback_context,
        )

        # Append peer review summary if provided
        if review_result:
            sep = "=" * 70
            paper += f"\n\n{sep}\nPEER REVIEW SUMMARY\n{sep}\n"
            paper += review_result.get("summary", "")

        return paper

    def _stat_to_results(self, stat: dict) -> Dict:
        """AnalysisResult.to_dict() → write_full_paper의 results 포맷으로 변환."""
        paper_summary = stat.get("paper_summary", "")
        model_vars = stat.get("model_vars", [])

        main_findings = []
        main_findings.append(
            f"총 분석 대상 {stat.get('n_total', 0):,}명 중 "
            f"{stat.get('outcome_label', stat.get('outcome', ''))} 경험자 "
            f"{stat.get('n_outcome', 0):,}명 ({stat.get('outcome_rate', 0.0):.1f}%)"
        )
        for v in model_vars:
            if v.get("significant"):
                main_findings.append(
                    f"{v.get('label', v.get('variable', ''))} — "
                    f"adjusted OR {v.get('or_formatted', '')}, {v.get('p_formatted', '')}"
                )

        metrics = stat.get("model_metrics", {})
        pseudo_r2 = metrics.get("pseudo_r2")
        meta = f"Nagelkerke pseudo-R²={pseudo_r2:.3f}" if pseudo_r2 else ""

        return {
            "summary": paper_summary,
            "main_findings": main_findings or [paper_summary],
            "descriptive": str(stat.get("descriptive_stats", "")),
            "subgroup": str(stat.get("subgroup_results", "Not applicable")),
            "sensitivity": meta,
            "conclusion": f"본 연구는 {stat.get('outcome_label', '')}의 독립적 관련 요인을 규명하였다.",
        }

    # ------------------------------------------------------------------
    # DOCX export
    # ------------------------------------------------------------------

    def export_docx(self, paper_text: str, output_path: str | Path, title: str = "Research Paper") -> Path:
        """논문 텍스트를 Word DOCX로 저장.

        python-docx가 없으면 graceful fallback (.txt 저장).
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            t = doc.add_heading(title, level=0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Author
            author_p = doc.add_paragraph(self._profile.author_name)
            author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_p.runs[0].italic = True

            doc.add_paragraph()

            current_section = None
            section_pattern = re.compile(
                r"^={10,}$|^(ABSTRACT|INTRODUCTION|METHODS|RESULTS|DISCUSSION|PEER REVIEW SUMMARY)$",
                re.IGNORECASE,
            )

            for line in paper_text.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if section_pattern.match(stripped) and not stripped.startswith("="):
                    current_section = stripped
                    doc.add_heading(stripped.title(), level=1)
                elif stripped.startswith("="):
                    continue
                elif stripped.startswith("By "):
                    continue
                else:
                    p = doc.add_paragraph(stripped)
                    p.runs[0].font.size = Pt(11) if p.runs else None

            docx_path = output_path.with_suffix(".docx")
            doc.save(str(docx_path))
            _log.info("DOCX saved: %s", docx_path)
            return docx_path

        except ImportError:
            _log.warning("python-docx not installed — saving as .txt fallback")
            txt_path = output_path.with_suffix(".txt")
            txt_path.write_text(paper_text, encoding="utf-8")
            return txt_path
        except Exception as e:
            _log.error("DOCX export failed: %s", e)
            txt_path = output_path.with_suffix(".txt")
            txt_path.write_text(paper_text, encoding="utf-8")
            return txt_path

    # ------------------------------------------------------------------
    # Internal
    # ------------------------------------------------------------------

    def _build_exemplars_block(self, *, n: int = 2, per_chars: int = 900) -> str:
        """yoosun_cho.json raw_examples를 few-shot으로 박는 블록.

        ★ FIX 3 (이전 1,500자×3=4,500자→ 900자×2=1,800자):
            user_prompt 앞에 4,500자 박으면 LLM 응답 토큰을 압박해 Methods 862자처럼
            짧아짐. 2개×900자 = 1,800자로 축소 — mimic 신호는 충분히 살아 있음.
        """
        try:
            import json as _j
            from pathlib import Path as _P
            p = _P("data/author_profiles/yoosun_cho.json")
            if not p.exists():
                return ""
            prof = _j.loads(p.read_text(encoding="utf-8"))
            exs = prof.get("raw_examples") or []
            if not exs:
                return ""
            picks = exs[:n]
            block = ("## AUTHOR'S ACTUAL WRITING (verbatim — mimic the cadence)\n\n")
            for i, ex in enumerate(picks, 1):
                block += f"### Ex{i}\n{str(ex)[:per_chars]}\n\n"
            block += ("Mimic the above: sentence rhythm, verb selection, hedging, "
                       "topic-sentence structure. Do NOT default to generic LLM tone.\n\n")
            return block
        except Exception:
            return ""

    def _generate(self, system_prompt: str, user_prompt: str,
                   *, max_tokens: int = 8192) -> str:
        # ★ FIX 1 (라인 786 원본): max_tokens 미지정으로 Methods/Results가
        #   기본 2048 토큰(~1,500자)에서 잘리는 사고. 섹션별로 명시 전달.
        # ★ FIX 4: anti_meta 룰 압축 — 너무 길면 LLM이 본 지시를 덮어버림.
        # ★ FIX 13 (자산화 1단계 연결): humanize 카탈로그를 system_prompt에 주입.
        #   12,301편에서 추출한 sentence-level humanization signal 7종 × 1 예문.
        humanize_block = ""
        try:
            from src.knowledge.humanize_extractor import get_humanize_block
            humanize_block = get_humanize_block(sample_per_kind=1)
        except Exception:
            pass

        anti_meta = (
            "\n\n# OUTPUT RULES\n"
            "- Output ONLY the section content. No preamble, no meta.\n"
            "- Forbidden openings: 'This is', 'Let me', 'Let's', \"Here's\", 'Here is', "
            "'I will', \"I'll\", 'Sure', 'Certainly', 'Below is', 'Following is', "
            "'In this abstract/section/paper'.\n"
            "- Do NOT mention the author by name in meta ('mimicking Dr. Cho').\n"
            "- Do NOT include section header lines ('**Abstract**:', '# Methods:').\n"
            "- Start directly with the first sentence of the academic content.\n"
        )
        full_system = system_prompt + (("\n\n" + humanize_block) if humanize_block else "") + anti_meta

        out = self._client.generate(
            user_prompt,
            system_prompt=full_system,
            max_tokens=max_tokens,
        )
        # ★ Response sanitizer — 그래도 새어 나온 메타 코멘트 절단
        out = _strip_llm_meta(out)

        # ★ FIX 14 (자산화 2단계 연결): anti-AI 필터 — LLM 흔적 자동 정리
        try:
            from src.safety.anti_ai_filter import filter_text as _ai_filter, ai_score
            if out and len(out) > 200:
                before = ai_score(out)
                if before.score > 25:   # 명백한 AI-스러움일 때만 정리
                    cleaned, _b, after = _ai_filter(out, mode="gentle")
                    _log.info("[PaperWriter] anti-AI: score %.1f → %.1f (saved %d chars)",
                               before.score, after.score, len(out) - len(cleaned))
                    out = cleaned
        except Exception as _e:
            _log.warning("[PaperWriter] anti-AI filter 실패: %s", _e)

        # Safety gate — 임상 키워드 자동 격리 큐
        try:
            from src.safety import check_all
            if out:
                check_all(out, scope="paper_writer._generate", design="cross_sectional")
        except Exception:
            pass
        try:
            from src.safety.physician_review import review_required, queue_for_review
            if out:
                needed, _triggers = review_required(out)
                if needed:
                    queue_for_review(content=out[:8000], source="paper_writer._generate")
        except Exception:
            pass
        return out


def _strip_llm_meta(text: str) -> str:
    """LLM 응답에서 메타 코멘트(자기소개·서문) 제거.

    2026-05-31 사고: "This is an interesting intersection of topics – ...
    Let's structure an abstract that captures the essence of the study,
    mimicking the precise style and rigor of Dr. Yoosun Cho." 같은 메타가
    본문 첫 부분에 박혀 docx에 그대로 들어간 사고.

    가장 흔한 LLM 메타 패턴을 정규식으로 잡아 첫 1-2 문장 절단.
    """
    if not text:
        return text
    import re as _re_meta
    cleaned = text.lstrip()
    # 흔한 메타 패턴 — 본문 시작 전 1-2 문장
    meta_patterns = [
        # "This is an interesting/important/fascinating/complex intersection of topics ..."
        r"^This is (?:an? |the )?(?:interesting|important|fascinating|complex|"
        r"nuanced|fascinating|excellent|great|wonderful|critical)[^.!?]*[.!?]",
        # "Let's structure / Let me write / Let me craft / Let me draft"
        r"^Let'?s [^.!?]*[.!?]",
        r"^Let me [^.!?]*[.!?]",
        # "Here's / Here is the abstract/methods/section"
        r"^Here'?s (?:the |an? )?[^.!?]*[.!?]",
        r"^Here is (?:the |an? )?[^.!?]*[.!?]",
        # "I'll / I will / I'm going to write/draft/craft/structure"
        r"^I'?ll [^.!?]*[.!?]",
        r"^I will [^.!?]*[.!?]",
        r"^I'?m going to [^.!?]*[.!?]",
        # "Below is / Following is the abstract"
        r"^Below is [^.!?]*[.!?]",
        r"^Following is [^.!?]*[.!?]",
        # "In this abstract/section/paper, I will ..."
        r"^In this (?:abstract|section|paper|response)[^.!?]*[.!?]",
        # "Sure! Certainly! Absolutely! Of course!"
        r"^(?:Sure|Certainly|Absolutely|Of course)[,.!:][^.!?]*[.!?]",
        # "I'll draft an abstract that mimics..."
        r"^[^.!?]*(?:mimicking|mimic|emulating|matching|following) the [^.!?]*style[^.!?]*[.!?]",
        # "The following abstract / The abstract below"
        r"^The following [^.!?]*[.!?]",
        # **Abstract** / # Abstract / Abstract: (섹션 라벨이 본문 앞에 박힘)
        r"^\*\*\s*(?:Abstract|Introduction|Methods|Results|Discussion)\s*\*\*[:\s]*",
        r"^#+\s*(?:Abstract|Introduction|Methods|Results|Discussion)[:\s]*\n",
        r"^(?:Abstract|Introduction|Methods|Results|Discussion):\s*\n",
    ]
    # ★ FIX 4a: 5번 → 2번. 5번 반복은 본문 첫 단락 2-3개가 통째로 삭제될 위험.
    for _ in range(2):
        matched = False
        for pat in meta_patterns:
            m = _re_meta.match(pat, cleaned, _re_meta.IGNORECASE | _re_meta.DOTALL)
            if m:
                cleaned = cleaned[m.end():].lstrip()
                matched = True
                break
        if not matched:
            break

    # ★ FIX 8 (Results 첫 줄 '### Characteristics...' 누출 차단):
    #   본문 전체의 줄 시작 markdown 헤딩 마커 #~###### 를 평문 단락 제목으로 평탄화.
    #   docx 변환 시 '#'가 본문에 그대로 들어가 깨지는 것 차단. 헤딩 텍스트 자체는 보존.
    cleaned = _re_meta.sub(r"^#{1,6}[ \t]+", "", cleaned, flags=_re_meta.MULTILINE)
    #   bold sub-header (**Characteristics**:) → 평문화
    cleaned = _re_meta.sub(r"^\*\*([^*\n]{2,80})\*\*[ \t]*:?[ \t]*$",
                            r"\1", cleaned, flags=_re_meta.MULTILINE)

    # ★ FIX 4b: 2차 안전망 — 키워드를 "단락 시작 prefix"로만 좁혀 본문 보호.
    #   원본은 "let's" 같은 일반 단어가 본문 중간에 있어도 단락 통째 삭제 →
    #   첫 30자 이내에 명시적 메타 prefix가 있을 때만 삭제. 1회만.
    paragraphs = cleaned.split("\n\n", 1)
    if len(paragraphs) == 2:
        first_head = paragraphs[0][:60].lower().lstrip()
        meta_prefixes = (
            "let's ", "let me ", "this is an interesting", "this is a fascinating",
            "this is the abstract", "this is the introduction",
            "i'll write", "i will write", "i'll draft", "i'll structure",
            "i'll capture", "i'll craft", "here's the abstract", "here is the abstract",
            "here's the introduction", "here is the introduction",
            "here's the methods", "here is the methods",
            "here's the results", "here is the results",
            "here's the discussion", "here is the discussion",
            "below is the", "following is the",
            "mimicking dr.", "emulating dr.", "in dr. cho's style",
            "i'm going to ", "i am going to ",
        )
        if first_head.startswith(meta_prefixes):
            cleaned = paragraphs[1].lstrip()
    return cleaned


# (_do_generate 분리는 제거 — _generate 메소드가 직접 본문을 가짐)
