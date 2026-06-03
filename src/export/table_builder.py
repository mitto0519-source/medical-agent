"""Table builder — 통계결과 → python-docx Table (학술지 수준 가독성).

양식 (`data/templates/manuscript_template.json` 의 tables 섹션):
  · 세 줄 표(Three-line table) — Top/Header-bottom/Bottom horizontal lines only,
    내부 세로/가로 격자 없음. NEJM/Lancet/JAMA 표준.
  · Times New Roman 9pt, line_spacing single
  · 헤더 굵게, 배경색 없음 (학술지 흰색)
  · 숫자열 right-align, 카테고리열 left-align
  · p-value: `P < 0.001` / `P = 0.026` (P italic)
  · 캡션은 표 위에, `Table N.` 굵게 + " {제목}"
  · footnote italic 8pt

dispatcher: `render_publication_table(doc, tbl_dict, tpl)` —
  `tbl_dict["type"]` ∈ {"baseline", "regression", "cross", "raw"}.
"""
from __future__ import annotations

from typing import Any, Dict, List, Optional, Tuple

from src.config.logging_config import get_logger

_log = get_logger(__name__)

# ── 학술지  색·간격 (template 미주입 시 fallback) ────────────────────────
_BORDER_BLACK = "000000"
_TABLE_FONT = "Times New Roman"
_TABLE_PT = 9
_CAPTION_PT = 10
_FOOTNOTE_PT = 8


# ── 셀/표 스타일 helpers ─────────────────────────────────────────────────────

def _qn(tag: str):
    from docx.oxml.ns import qn
    return qn(tag)


def _OxmlElement(name: str):
    from docx.oxml import OxmlElement
    return OxmlElement(name)


def _set_cell_borders(cell, *, top: Optional[float] = None,
                      bottom: Optional[float] = None,
                      left: Optional[float] = None,
                      right: Optional[float] = None):
    """셀별 horizontal border 두께 설정 (학술지 세 줄 표 구현). 단위: pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(_qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = _OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side, w in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        existing = tcBorders.find(_qn(f"w:{side}"))
        if existing is not None:
            tcBorders.remove(existing)
        b = _OxmlElement(f"w:{side}")
        if w is None or w <= 0:
            b.set(_qn("w:val"), "nil")
        else:
            b.set(_qn("w:val"), "single")
            b.set(_qn("w:sz"), str(int(w * 8)))  # docx unit = 1/8 pt
            b.set(_qn("w:color"), _BORDER_BLACK)
        tcBorders.append(b)


def _apply_three_line(table, header_rows: int = 1, top_pt: float = 1.25,
                       mid_pt: float = 0.75, bottom_pt: float = 1.25):
    """세 줄 표 — 표 외곽 top/bottom 굵게, 헤더 아래 가는 선, 내부 다 nil."""
    n_rows = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            top = top_pt if r_idx == 0 else None
            if r_idx == header_rows - 1:
                bottom = mid_pt
            elif r_idx == n_rows - 1:
                bottom = bottom_pt
            else:
                bottom = 0
            _set_cell_borders(cell, top=top, bottom=bottom, left=0, right=0)


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False,
                   alignment: str = "left", font: str = _TABLE_FONT,
                   size_pt: int = _TABLE_PT):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    cell.text = ""
    p = cell.paragraphs[0]
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                 "center": WD_ALIGN_PARAGRAPH.CENTER}
    p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)


def _set_cell_rich(cell, segments: List[Tuple[str, Dict]], *, alignment: str = "left",
                   font: str = _TABLE_FONT, size_pt: int = _TABLE_PT):
    """셀 안에 여러 서식 segment 삽입. segments=[(text, {bold,italic}), ...]"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    cell.text = ""
    p = cell.paragraphs[0]
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                 "center": WD_ALIGN_PARAGRAPH.CENTER}
    p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    for txt, fmt in segments:
        run = p.add_run(txt)
        run.bold = fmt.get("bold", False)
        run.italic = fmt.get("italic", False)
        run.font.name = font
        run.font.size = Pt(size_pt)


def _add_caption(doc, label: str, text: str, *, position: str = "above",
                 size_pt: int = _CAPTION_PT, font: str = _TABLE_FONT,
                 alignment: str = "left"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if alignment == "left" else WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6) if position == "above" else Pt(2)
    pf.space_after = Pt(4) if position == "above" else Pt(8)
    pf.line_spacing = 1.0
    if label:
        run = p.add_run(label + " ")
        run.bold = True
        run.font.name = font
        run.font.size = Pt(size_pt)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)


def _add_footnote(doc, text: str, *, size_pt: int = _FOOTNOTE_PT, font: str = _TABLE_FONT):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    run.italic = True
    run.font.name = font
    run.font.size = Pt(size_pt)


# ── p-value / OR(CI) 양식 ────────────────────────────────────────────────────

def _fmt_p(p: Optional[float]) -> List[Tuple[str, dict]]:
    """`P < 0.001` / `P = 0.026` (P italic). 셀에 넣을 segments 반환."""
    if p is None:
        return [("", {})]
    if p < 0.001:
        return [("P", {"italic": True}), (" < 0.001", {})]
    return [("P", {"italic": True}), (f" = {p:.3f}", {})]


def _fmt_or_ci(or_val: Optional[float], lo: Optional[float], hi: Optional[float]) -> str:
    if or_val is None:
        return ""
    if lo is not None and hi is not None:
        return f"{or_val:.2f} ({lo:.2f}–{hi:.2f})"
    return f"{or_val:.2f}"


def _fmt_n_pct(n: Optional[int], pct: Optional[float]) -> str:
    if n is None and pct is None:
        return ""
    if pct is None:
        return f"{n:,}"
    if n is None:
        return f"{pct:.1f}"
    return f"{n:,} ({pct:.1f})"


def _fmt_mean_sd(mean: Optional[float], sd: Optional[float]) -> str:
    if mean is None:
        return ""
    if sd is None:
        return f"{mean:.2f}"
    return f"{mean:.2f} ± {sd:.2f}"


# ── 학술지  표 빌더 (재작성, 시그니처 호환) ──────────────────────────────

def baseline_characteristics_table(doc, data: List[Dict],
                                    caption: str = "Table 1. Baseline Characteristics"):
    """기준 특성 표 (학술지 세 줄 표).

    data: [{"variable": str, "n": int, "pct": float, "mean": float, "sd": float}, ...]
    caption: "Table 1. ..."로 시작하면 label 분리.
    """
    label, text = _split_caption(caption, default_label="Table 1.")
    _add_caption(doc, label, text, position="above")

    headers = ["Characteristic", "n (%) or Mean ± SD"]
    table = doc.add_table(rows=1 + len(data), cols=len(headers))

    # 헤더
    for i, h in enumerate(headers):
        align = "left" if i == 0 else "right"
        _set_cell_text(table.cell(0, i), h, bold=True, alignment=align)

    # 데이터 행
    for r_idx, row_data in enumerate(data, start=1):
        var = str(row_data.get("variable", ""))
        # value 합성: mean±SD 있으면 그것, 없으면 n (%)
        if row_data.get("mean") is not None:
            value = _fmt_mean_sd(row_data["mean"], row_data.get("sd"))
        else:
            value = _fmt_n_pct(row_data.get("n"), row_data.get("pct"))
        # 들여쓰기 카테고리 ("  — value")는 그대로 유지
        _set_cell_text(table.cell(r_idx, 0), var, alignment="left")
        _set_cell_text(table.cell(r_idx, 1), value, alignment="right")

    _apply_three_line(table)
    _add_footnote(doc, "Values are n (%) for categorical variables or mean ± SD for continuous variables.")
    return table


def regression_table(doc, results: List[Dict],
                      caption: str = "Table 2. Adjusted Odds Ratios"):
    """로지스틱 회귀 결과 표 (학술지 세 줄 표, p-value italic P).

    results: [{"variable": str, "or": float, "ci_low": float, "ci_high": float,
               "p_value": float}, ...]  ('or' 없으면 'beta' 사용)
    """
    label, text = _split_caption(caption, default_label="Table 2.")
    _add_caption(doc, label, text, position="above")

    has_or = any(r.get("or") is not None for r in results)
    headers = ["Variable", "aOR (95% CI)" if has_or else "β (95% CI)", "P-value"]
    table = doc.add_table(rows=1 + len(results), cols=len(headers))

    for i, h in enumerate(headers):
        align = "left" if i == 0 else "right"
        # P-value 헤더는 italic P
        if h == "P-value":
            _set_cell_rich(table.cell(0, i),
                            [("P", {"bold": True, "italic": True}),
                             ("-value", {"bold": True})], alignment="right")
        else:
            _set_cell_text(table.cell(0, i), h, bold=True, alignment=align)

    for r_idx, r in enumerate(results, start=1):
        var = str(r.get("variable", "") or r.get("label", ""))
        if has_or:
            val = _fmt_or_ci(r.get("or"), r.get("ci_low"), r.get("ci_high"))
        else:
            b = r.get("beta")
            ci = f" ({r.get('ci_low', 0):.2f}–{r.get('ci_high', 0):.2f})" if r.get("ci_low") is not None else ""
            val = f"{b:.3f}{ci}" if b is not None else ""
        _set_cell_text(table.cell(r_idx, 0), var, alignment="left")
        _set_cell_text(table.cell(r_idx, 1), val, alignment="right")
        _set_cell_rich(table.cell(r_idx, 2), _fmt_p(r.get("p_value")), alignment="right")

    _apply_three_line(table)
    _add_footnote(doc,
        "aOR, adjusted odds ratio; CI, confidence interval. "
        "P-values from survey-weighted logistic regression.")
    return table


def cross_table(doc, contingency: List[List], row_labels: List[str],
                col_labels: List[str], caption: str = "Table 3. Cross Tabulation"):
    """교차표 (chi-square 결과). 학술지 세 줄 표."""
    label, text = _split_caption(caption, default_label="Table 3.")
    _add_caption(doc, label, text, position="above")

    table = doc.add_table(rows=1 + len(row_labels), cols=1 + len(col_labels))

    _set_cell_text(table.cell(0, 0), "", bold=True, alignment="left")
    for j, cl in enumerate(col_labels, start=1):
        _set_cell_text(table.cell(0, j), cl, bold=True, alignment="right")

    for i, rl in enumerate(row_labels, start=1):
        _set_cell_text(table.cell(i, 0), rl, alignment="left")
        for j, val in enumerate(contingency[i - 1], start=1):
            _set_cell_text(table.cell(i, j), str(val), alignment="right")

    _apply_three_line(table)
    return table


def raw_table(doc, headers: List[str], rows: List[List[str]],
              caption: str = "Table",
              footnote: str = "",
              numeric_cols: Optional[List[int]] = None):
    """임의 raw 표 (헤더 + N개 행). word_exporter publication dispatcher가 활용."""
    label, text = _split_caption(caption, default_label="Table.")
    _add_caption(doc, label, text, position="above")

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    numeric_cols = set(numeric_cols or [])

    for i, h in enumerate(headers):
        align = "right" if i in numeric_cols else "left"
        _set_cell_text(table.cell(0, i), h, bold=True, alignment=align)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            align = "right" if c_idx in numeric_cols else "left"
            _set_cell_text(table.cell(r_idx, c_idx), str(val), alignment=align)

    _apply_three_line(table)
    if footnote:
        _add_footnote(doc, footnote)
    return table


# ── Publication dispatcher (word_exporter가 호출) ────────────────────────────

def render_publication_table(doc, tbl: Dict, tpl: Optional[Dict] = None):
    """`tbl["type"]`에 따라 위 빌더로 분기. word_exporter._embed_table()이 호출."""
    ttype = tbl.get("type", "raw")
    caption = tbl.get("caption", "")
    n = tbl.get("n")
    if n and not caption.lower().startswith(f"table {n}"):
        caption = f"Table {n}. {caption}"

    if ttype == "baseline":
        return baseline_characteristics_table(doc, tbl.get("data", []), caption=caption)
    if ttype == "regression":
        return regression_table(doc, tbl.get("data", []), caption=caption)
    if ttype == "cross":
        return cross_table(doc, tbl.get("data", []),
                            tbl.get("row_labels", []),
                            tbl.get("col_labels", []), caption=caption)
    # raw
    return raw_table(doc, tbl.get("headers", []), tbl.get("rows", []),
                      caption=caption, footnote=tbl.get("footnote", ""),
                      numeric_cols=tbl.get("numeric_cols"))


def _split_caption(caption: str, default_label: str = "Table.") -> Tuple[str, str]:
    """'Table 1. Foo' → ('Table 1.', 'Foo')."""
    import re
    m = re.match(r"^(Table\s+\d+\.?)\s*(.*)$", caption.strip())
    if m:
        return m.group(1), m.group(2).strip()
    return default_label, caption


# ───────────────────────────────────────────────────────────────────────────
# StatBridge stat_result → Table 1/2 (호환 유지)
# ───────────────────────────────────────────────────────────────────────────

_VAR_LABELS: Dict[str, str] = {
    "sex": "Sex, female", "sex_2": "Sex, female",
    "sleep_hours": "Sleep duration (h/d)", "screen_time": "Screen time (h/d)",
    "smoking": "Current smoking", "alcohol": "Alcohol use",
    "physical_act": "Physical activity (d/wk)", "bmi": "BMI (kg/m²)",
    "grade": "School grade", "family_econ": "Socioeconomic status (low)",
    "academic_perf": "Academic performance (low)", "stress": "Perceived stress (high)",
    "depression": "Depressive mood", "suicidal": "Suicidal ideation",
    "obesity": "Obesity (BMI ≥ 25)", "hypertension": "Hypertension",
    "diabetes": "Diabetes", "metabolic_syn": "Metabolic syndrome",
    "zcb_freq": "Zero-calorie beverage (per 1-level)",
    "ssb_freq": "Sugar-sweetened beverage (per 1-level)",
    "caffeine_freq": "Caffeine beverage (per 1-level)",
    "breakfast": "Breakfast skipping",
    "school_type": "School type (high school)",
}


def _prettify_var(col: str) -> str:
    if col in _VAR_LABELS:
        return _VAR_LABELS[col]
    return col.replace("_", " ").title()


def stat_result_to_table1_markdown(stat_result: dict) -> str:
    """StatBridge stat_result → Table 1 마크다운 문자열 (호환 유지)."""
    n_total = stat_result.get("n_total", 0)
    desc = stat_result.get("descriptive_stats", {})
    outcome_lbl = stat_result.get("outcome_label", stat_result.get("outcome", "Outcome"))

    lines = [
        f"**Table 1. Characteristics of Study Participants (N = {n_total:,})**",
        "",
        "| Variable | Value |",
        "|----------|-------|",
        f"| Total N | {n_total:,} |",
        f"| {outcome_lbl}, n (%) | "
        f"{stat_result.get('n_outcome', 0):,} ({stat_result.get('outcome_rate', 0.0):.1f}%) |",
    ]
    for col, stats in desc.items():
        label = _prettify_var(col)
        if "mean" in stats:
            m, s = stats["mean"], stats["std"]
            lines.append(f"| {label}, mean ± SD | {m:.2f} ± {s:.2f} |")
        elif "categories" in stats:
            first = True
            for cat_val, pct in stats["categories"].items():
                if first:
                    lines.append(f"| {label} |  |")
                    first = False
                lines.append(f"|   — {cat_val} | {pct * 100:.1f}% |")
    return "\n".join(lines)


def stat_result_to_table2_markdown(stat_result: dict) -> str:
    """StatBridge stat_result → Table 2 마크다운 문자열 (호환 유지)."""
    outcome_lbl = stat_result.get("outcome_label", stat_result.get("outcome", "Outcome"))
    model_vars = stat_result.get("model_vars", [])
    lines = [
        f"**Table 2. Logistic Regression Results — Odds Ratios for {outcome_lbl}**",
        "",
        "| Variable | aOR | 95% CI | P-value |",
        "|----------|-----|--------|---------|",
    ]
    for v in model_vars:
        label = v.get("label") or _prettify_var(v.get("variable", ""))
        or_f = v.get("or_formatted", "") or (f"{v.get('or_value'):.2f}" if v.get("or_value") else "")
        ci_f = v.get("ci_formatted", "")
        if not ci_f and v.get("ci_lower") is not None:
            ci_f = f"{v['ci_lower']:.2f}–{v['ci_upper']:.2f}"
        p = v.get("p_value")
        p_f = "<0.001" if (p is not None and p < 0.001) else (f"{p:.3f}" if p is not None else "")
        sig = " *" if v.get("significant") else ""
        lines.append(f"| {label}{sig} | {or_f} | {ci_f} | {p_f} |")
    lines += ["", "*P < 0.05; aOR, adjusted odds ratio; CI, confidence interval"]
    return "\n".join(lines)


def stat_result_to_tables_docx_bytes(stat_result: dict) -> bytes:
    """StatBridge stat_result → 학술지 세 줄 표 DOCX 바이트 반환."""
    import io
    try:
        from docx import Document
    except ImportError:
        _log.error("python-docx 미설치")
        return b""

    doc = Document()
    n_total = stat_result.get("n_total", 0)
    outcome_lbl = stat_result.get("outcome_label", stat_result.get("outcome", "Outcome"))
    desc = stat_result.get("descriptive_stats", {})
    outcome_key = stat_result.get("outcome", "")

    # Table 1 데이터
    table1_data = [
        {"variable": "Total N", "n": n_total, "pct": None, "mean": None, "sd": None},
        {"variable": outcome_lbl,
         "n": stat_result.get("n_outcome", 0),
         "pct": stat_result.get("outcome_rate", 0.0),
         "mean": None, "sd": None},
    ]
    for col, stats in desc.items():
        if col == outcome_key:
            continue
        label = _prettify_var(col)
        if "mean" in stats:
            table1_data.append({"variable": label, "n": n_total,
                                 "pct": None, "mean": stats["mean"], "sd": stats["std"]})
        elif "categories" in stats:
            for cat_val, pct in list(stats["categories"].items())[:3]:
                table1_data.append({"variable": f"  — {cat_val}",
                                     "n": round(pct * n_total),
                                     "pct": pct * 100, "mean": None, "sd": None})

    baseline_characteristics_table(
        doc, table1_data,
        caption=f"Table 1. Characteristics of Study Participants (N = {n_total:,})"
    )

    # Table 2 데이터
    model_vars = stat_result.get("model_vars", [])
    table2_data = []
    for v in model_vars:
        label = v.get("label") or _prettify_var(v.get("variable", ""))
        table2_data.append({
            "variable": label,
            "or": v.get("or_value"),
            "ci_low": v.get("ci_lower"),
            "ci_high": v.get("ci_upper"),
            "p_value": v.get("p_value"),
        })
    if table2_data:
        regression_table(
            doc, table2_data,
            caption=f"Table 2. Adjusted Odds Ratios for {outcome_lbl}"
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
