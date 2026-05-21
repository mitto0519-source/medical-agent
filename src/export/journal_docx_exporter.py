"""JournalDocxExporter — 저널 스타일이 완전 적용된 DOCX 생성.

JournalStyle(journal_registry.py)의 서식 규칙을 python-docx로 정확히 구현:
  - 폰트, 글자 크기, 줄간격, 여백
  - 섹션 순서 (저널마다 다름: IJERPH는 "Materials and Methods" + "Conclusions")
  - 초록 구조 (저널마다 다름: BMJ Open은 7개 소제목)
  - 참고문헌 스타일 (Vancouver / APA)
  - 키워드, 제목 자수 제한
  - EndNote XML / BibTeX 동시 생성
"""
from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Dict, List, Optional

from src.config.logging_config import get_logger
from src.export.journal_registry import JournalStyle, get_registry
from src.export.reference_library import Reference, ReferenceLibrary, format_reference

_log = get_logger(__name__)

_OUTPUT_DIR = Path("data/drafts/word")
_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


class JournalDocxExporter:
    """저널 서식에 맞춰 논문을 Word .docx로 출력.

    Usage:
        exporter = JournalDocxExporter("jkms")
        path = exporter.export(
            title="청소년 스마트폰 과사용과 수면 부족",
            sections={"Abstract": "...", "Introduction": "...", ...},
            references=ref_library,   # ReferenceLibrary 또는 List[Reference]
            authors=["Yoosun Cho"],
            keywords=["adolescent", "smartphone", "sleep"],
            output_path="data/drafts/word/paper.docx",
        )
        # 동시에 EndNote XML + BibTeX 도 저장됨
    """

    def __init__(self, journal: str | JournalStyle = "jkms"):
        if isinstance(journal, JournalStyle):
            self.style = journal
        else:
            self.style = get_registry().get(journal)
        _log.info("JournalDocxExporter 초기화: %s (%s)", self.style.name, self.style.id)

    # ── 공개 API ─────────────────────────────────────────────────────────────

    def export(
        self,
        title: str,
        sections: Dict[str, str],
        references: Optional[ReferenceLibrary | List[Reference]] = None,
        authors: Optional[List[str]] = None,
        affiliation: str = "",
        keywords: Optional[List[str]] = None,
        running_title: str = "",
        output_path: Optional[str] = None,
        also_save_endnote: bool = True,
        also_save_bibtex: bool = True,
    ) -> str:
        """저널 서식 DOCX 생성 후 파일 경로 반환."""
        try:
            from docx import Document
        except ImportError:
            _log.error("python-docx 미설치 — pip install python-docx")
            raise

        doc = Document()
        self._apply_page_setup(doc)
        self._apply_default_style(doc)

        # ── 표지 영역 ──────────────────────────────────────────────────────
        self._add_title(doc, title, authors or ["Yoosun Cho"], affiliation, running_title)
        if keywords:
            self._add_keywords(doc, keywords)

        # ── 섹션 본문 ──────────────────────────────────────────────────────
        for sec_name in self.style.section_order:
            # "References"는 마지막에 별도 처리
            if sec_name.lower() == "references":
                continue
            # 저널별 섹션 이름 매핑 (IJERPH: Methods → Materials and Methods)
            text = self._find_section_text(sections, sec_name)
            if not text:
                continue
            self._add_section_heading(doc, sec_name)
            if sec_name.lower() == "abstract":
                self._add_structured_abstract(doc, text)
            else:
                self._add_body_paragraphs(doc, text)

        # ── 참고문헌 ───────────────────────────────────────────────────────
        refs = self._resolve_refs(references)
        if refs:
            self._add_section_heading(doc, "References")
            for i, r in enumerate(refs, 1):
                self._add_reference_entry(doc, r, i)

        # ── 저장 ───────────────────────────────────────────────────────────
        if not output_path:
            safe = re.sub(r"[^\w\s\-]", "_", title)[:60].strip()
            output_path = str(_OUTPUT_DIR / f"{safe}_{self.style.id}.docx")

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        _log.info("DOCX 저장 완료: %s", output_path)

        # ── EndNoteXML / BibTeX 동시 저장 ─────────────────────────────────
        if refs and (also_save_endnote or also_save_bibtex):
            slug = re.sub(r"[^\w]", "_", title)[:60]
            ref_lib = ReferenceLibrary(slug)
            for r in refs:
                ref_lib.add_manual(r)
            if also_save_endnote:
                ref_lib.save_endnote_xml()
            if also_save_bibtex:
                ref_lib.save_bibtex()
            ref_lib.save()

        return output_path

    def export_bytes(self, **kwargs) -> bytes:
        """Streamlit 다운로드용 — 파일 저장 없이 바이트 반환."""
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        kwargs.setdefault("also_save_endnote", False)
        kwargs.setdefault("also_save_bibtex", False)
        try:
            self.export(output_path=tmp_path, **kwargs)
            with open(tmp_path, "rb") as f:
                return f.read()
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    def get_endnote_xml_bytes(self, refs: List[Reference]) -> bytes:
        """EndNote XML 바이트 반환 (Streamlit 다운로드용)."""
        from src.export.reference_library import to_endnote_xml
        return to_endnote_xml(refs).encode("utf-8")

    def get_bibtex_bytes(self, refs: List[Reference]) -> bytes:
        """BibTeX 바이트 반환."""
        from src.export.reference_library import to_bibtex
        return to_bibtex(refs).encode("utf-8")

    # ── 페이지 설정 ───────────────────────────────────────────────────────────

    def _apply_page_setup(self, doc):
        from docx.shared import Cm
        m = self.style.margins_cm
        for section in doc.sections:
            section.top_margin = Cm(m["top"])
            section.bottom_margin = Cm(m["bottom"])
            section.left_margin = Cm(m["left"])
            section.right_margin = Cm(m["right"])

    def _apply_default_style(self, doc):
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import docx.oxml

        normal = doc.styles["Normal"]
        normal.font.name = self.style.font_name
        normal.font.size = Pt(self.style.font_size)

        # 줄간격 설정
        pf = normal.paragraph_format
        ls = self.style.line_spacing
        if ls == 2.0:
            from docx.enum.text import WD_LINE_SPACING
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        elif ls == 1.5:
            from docx.enum.text import WD_LINE_SPACING
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            from docx.shared import Pt as Pt2
            pf.line_spacing = Pt2(self.style.font_size * ls)

    # ── 제목/저자/키워드 ─────────────────────────────────────────────────────

    def _add_title(self, doc, title: str, authors: List[str], affiliation: str, running_title: str):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(self.style.font_size + 2)
        r.font.name = self.style.font_name

        if authors:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(", ".join(authors))
            r2.font.size = Pt(self.style.font_size)
            r2.font.name = self.style.font_name

        if affiliation:
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r3 = p3.add_run(affiliation)
            r3.italic = True
            r3.font.size = Pt(self.style.font_size - 1)

        # Running title (short title)
        max_chars = self.style.formatting.get("running_title_max_chars", 60)
        rt = running_title or (title[:max_chars] + ("..." if len(title) > max_chars else ""))
        if rt:
            p4 = doc.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r4 = p4.add_run(f"Running title: {rt}")
            r4.italic = True
            r4.font.size = Pt(self.style.font_size - 1)

        doc.add_paragraph()

    def _add_keywords(self, doc, keywords: List[str]):
        from docx.shared import Pt
        max_kw = self.style.formatting.get("keywords_max", 6)
        shown = keywords[:max_kw] if max_kw else keywords
        p = doc.add_paragraph()
        r = p.add_run("Keywords: ")
        r.bold = True
        r.font.name = self.style.font_name
        p.add_run("; ".join(shown)).font.name = self.style.font_name
        doc.add_paragraph()

    # ── 섹션 헤딩 ─────────────────────────────────────────────────────────────

    def _add_section_heading(self, doc, section_name: str):
        from docx.shared import Pt
        from docx.oxml.ns import qn

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(section_name.upper())
        r.bold = True
        r.font.size = Pt(self.style.font_size)
        r.font.name = self.style.font_name

    # ── 초록 구조화 ───────────────────────────────────────────────────────────

    def _add_structured_abstract(self, doc, abstract_text: str):
        """저널별 초록 소제목(Background/Methods/Results/Conclusion)을 적용."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        structure = self.style.abstract_structure
        # 소제목 기반 분할 시도
        pattern = "|".join(re.escape(s) for s in structure)
        parts = re.split(f"({pattern}):", abstract_text, flags=re.IGNORECASE)

        if len(parts) > 1:
            # 구조화된 초록 파싱 성공
            i = 0
            while i < len(parts):
                token = parts[i].strip()
                if any(token.lower() == s.lower() for s in structure):
                    heading = token
                    body = parts[i + 1].strip() if i + 1 < len(parts) else ""
                    p = doc.add_paragraph()
                    r_h = p.add_run(heading + ": ")
                    r_h.bold = True
                    r_h.font.name = self.style.font_name
                    p.add_run(body).font.name = self.style.font_name
                    i += 2
                else:
                    if token:
                        self._add_body_paragraphs(doc, token)
                    i += 1
        else:
            # 구조화 불가 → 단일 단락
            self._add_body_paragraphs(doc, abstract_text)

    # ── 본문 단락 ─────────────────────────────────────────────────────────────

    def _add_body_paragraphs(self, doc, text: str):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for para in text.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            # 소제목 감지 (짧고 줄바꿈 직후인 경우)
            lines = para.split("\n")
            if len(lines) > 1 and len(lines[0]) < 60 and lines[0].endswith(":"):
                p = doc.add_paragraph()
                r = p.add_run(lines[0])
                r.bold = True
                r.font.name = self.style.font_name
                rest = "\n".join(lines[1:]).strip()
                if rest:
                    p2 = doc.add_paragraph(rest)
                    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p2.runs:
                        run.font.name = self.style.font_name
            else:
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = self.style.font_name

    # ── 참고문헌 항목 ─────────────────────────────────────────────────────────

    def _add_reference_entry(self, doc, ref: Reference, index: int):
        from docx.shared import Pt, Cm
        formatted = format_reference(ref, self.style.reference_style, index)
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(formatted)
        r.font.size = Pt(self.style.font_size - 1)
        r.font.name = self.style.font_name

    # ── 헬퍼 ─────────────────────────────────────────────────────────────────

    def _find_section_text(self, sections: Dict[str, str], sec_name: str) -> str:
        """대소문자 무관 섹션 텍스트 찾기. IJERPH 'Materials and Methods' 등 처리."""
        # 직접 매칭
        if sec_name in sections:
            return sections[sec_name]
        # 대소문자 무관
        for k, v in sections.items():
            if k.lower() == sec_name.lower():
                return v
        # 부분 일치 (Materials and Methods ↔ Methods)
        for k, v in sections.items():
            if sec_name.lower() in k.lower() or k.lower() in sec_name.lower():
                return v
        return ""

    @staticmethod
    def _resolve_refs(references) -> List[Reference]:
        if references is None:
            return []
        if isinstance(references, ReferenceLibrary):
            return references.get_refs()
        if isinstance(references, list):
            return [r for r in references if isinstance(r, Reference)]
        return []
