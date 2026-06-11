"""Citation Workflow — 레퍼런스 차용 검수 + 본문 넘버링 인용 + Word/EndNote 풀셋.

사용자 목표:
  레퍼런스 N개 투입 → ① 차용 가능성 검수(내 논문에 쓸 만한지)
  → ② 본문에 EndNote식 넘버링 인용 [n] 정확한 위치 삽입
  → ③ 참고문헌 목록(번호순 Vancouver) → ④ Word + EndNote 파일 풀셋.

검수·배치는 **임베딩 유사도(LLM-무관, 쿼터 무관)**. 메타/포맷/EndNote는 reference_library 재사용(규칙10).
"""
from __future__ import annotations

import contextlib
import io
import os
import re
import warnings
from typing import Dict, List, Optional, Tuple

# HF/tqdm 진행바가 Streamlit ScriptRunner 스레드의 '닫힌 stderr'에 쓰면
# "I/O operation on closed file"로 죽는다 → 진행바/경고를 원천 차단 (Gemini 수정과 동일 계열).
os.environ.setdefault("HF_HUB_DISABLE_PROGRESS_BARS", "1")
os.environ.setdefault("TRANSFORMERS_NO_ADVISORY_WARNINGS", "1")
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

from src.config.logging_config import get_logger
from src.export.reference_library import (
    Reference, _fetch_pubmed_xml, _parse_pubmed_xml, search_pubmed,
    format_vancouver, to_endnote_xml, to_bibtex,
)

_log = get_logger(__name__)
_MODEL = None
_BODY_ORDER = ["introduction", "results", "discussion", "methods", "abstract"]


@contextlib.contextmanager
def _quiet():
    """모델 로드/encode 동안 stderr·stdout·경고를 안전 버퍼로 돌려 닫힌-파일 오류를 차단."""
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        with contextlib.redirect_stderr(io.StringIO()), contextlib.redirect_stdout(io.StringIO()):
            yield


def _embedder():
    global _MODEL
    if _MODEL is None:
        with _quiet():
            from sentence_transformers import SentenceTransformer
            from src.config.models import get_embedding_model
            _MODEL = SentenceTransformer(get_embedding_model())
    return _MODEL


def _emb(text):
    """단일 텍스트 임베딩 — 닫힌-stderr 안전 가드 포함."""
    with _quiet():
        return _embedder().encode(text)


def _cos(a, b) -> float:
    import numpy as np
    a, b = np.asarray(a), np.asarray(b)
    d = (np.linalg.norm(a) * np.linalg.norm(b)) or 1.0
    return float(a @ b / d)


# ── ① 입력 파싱 + 메타 해결 ────────────────────────────────────────────────
def parse_reference_input(text: str) -> List[str]:
    """붙여넣은 레퍼런스 목록 → 항목 리스트 (PMID/DOI/제목 혼재 허용)."""
    out = []
    for raw in re.split(r"[\n;]+", text or ""):
        s = raw.strip()
        # 머리 목록 마커만 제거(불릿 또는 "1." "2)" 형태) — DOI(10.xxxx)는 보존
        s = re.sub(r"^\s*(?:[\-•*·]+\s+|\d{1,3}[.)\]]\s+)", "", s).strip()
        if len(s) >= 4:
            out.append(s)
    return out


def resolve_references(entries: List[str]) -> List[Reference]:
    """PMID(숫자) → 일괄 fetch, DOI/제목 → PubMed 검색 best match. 못 찾으면 제목만."""
    pmids, queries = [], []
    for e in entries:
        if re.fullmatch(r"\d{6,9}", e):
            pmids.append(e)
        else:
            m = re.search(r"10\.\d{4,9}/\S+", e)
            queries.append(m.group(0) if m else e)
    refs: List[Reference] = []
    if pmids:
        try:
            refs += _parse_pubmed_xml(_fetch_pubmed_xml(pmids))
        except Exception as ex:
            _log.warning("PMID fetch 실패: %s", ex)
    for q in queries:
        try:
            ids = search_pubmed(q, max_results=1)
            if ids:
                refs += _parse_pubmed_xml(_fetch_pubmed_xml(ids))
            else:
                refs.append(Reference(title=q))
        except Exception:
            refs.append(Reference(title=q))
    return refs


# ── ② 차용 가능성 검수 (임베딩 유사도) ─────────────────────────────────────
def screen_applicability(refs: List[Reference], paper_text: str,
                         threshold: float = 0.30) -> List[Dict]:
    """각 ref가 내 논문에 차용 가능한지 → [{ref, score, usable, reason}] (점수순)."""
    if not refs:
        return []
    pv = _emb((paper_text or "")[:2000])
    out = []
    for r in refs:
        rt = f"{r.title}. {r.abstract}".strip()[:2000]
        sim = _cos(pv, _emb(rt)) if rt.strip(". ") else 0.0
        out.append({
            "ref": r, "score": round(sim, 3), "usable": sim >= threshold,
            "reason": "주제 관련성 높음 — 차용 가능" if sim >= threshold else "관련성 낮음 — 차용 부적합",
        })
    out.sort(key=lambda x: -x["score"])
    return out


# ── ③ 본문 넘버링 인용 배치 ────────────────────────────────────────────────
def place_citations(sections: Dict[str, str],
                    usable_refs: List[Reference]) -> Tuple[Dict[str, str], List[Reference]]:
    """각 ref를 가장 관련 깊은 문장에 [n] 삽입. 등장 순으로 번호. (본문, 번호순 ref) 반환."""
    if not usable_refs:
        return dict(sections), []
    sents = []  # (sec, sentence_text)
    for sec in _BODY_ORDER:
        for sent in re.split(r"(?<=[.。?!])\s+", sections.get(sec) or ""):
            if len(sent.strip()) > 30:
                sents.append((sec, sent.strip()))
    if not sents:
        return dict(sections), []
    svecs = [_emb(s[1][:400]) for s in sents]

    # 각 ref → 가장 유사한 문장 인덱스
    assign: Dict[int, List[int]] = {}
    for ri, r in enumerate(usable_refs):
        rt = (f"{r.title}. {r.abstract}".strip()[:400]) or r.title
        rv = _emb(rt)
        best = max(range(len(sents)), key=lambda j: _cos(rv, svecs[j]))
        assign.setdefault(best, []).append(ri)

    # 인용 번호 = 문장 등장 순서
    numbered: Dict[int, int] = {}
    ordered: List[Reference] = []
    for sj in sorted(assign.keys()):
        for ri in assign[sj]:
            numbered[ri] = len(ordered) + 1
            ordered.append(usable_refs[ri])

    # 문장 뒤에 [n,m] 삽입
    new = dict(sections)
    by_sec: Dict[str, List[Tuple[str, str]]] = {}
    for sj, ri_list in assign.items():
        sec, txt = sents[sj]
        nums = sorted(numbered[ri] for ri in ri_list)
        by_sec.setdefault(sec, []).append((txt, "[" + ", ".join(map(str, nums)) + "]"))
    for sec, items in by_sec.items():
        body = sections.get(sec) or ""
        for txt, marker in items:
            body = body.replace(txt, f"{txt.rstrip()} {marker}", 1)
        new[sec] = body
    return new, ordered


# ── Multi-style formatters — 저널마다 다른 reference style 지원 ─────────────
def _authors_short(authors: List[str], max_n: int = 6) -> str:
    if not authors:
        return ""
    if len(authors) <= max_n:
        return ", ".join(authors)
    return ", ".join(authors[:max_n]) + ", et al"


def _author_last_first(name: str) -> str:
    """'John Smith' -> 'Smith J'  /  'Smith J' -> 'Smith J'."""
    parts = name.strip().split()
    if len(parts) < 2:
        return name
    if len(parts[-1]) <= 3 and parts[-1].isupper():
        return name  # already 'Smith J' form
    last = parts[-1]
    initials = "".join(p[0].upper() for p in parts[:-1])
    return f"{last} {initials}"


def format_ama(ref: Reference, index: int = 1) -> str:
    """AMA Manual of Style 11th ed. — JAMA Network journals."""
    authors = _authors_short([_author_last_first(a) for a in ref.authors], max_n=6)
    bits = []
    if authors: bits.append(authors + ".")
    if ref.title: bits.append(ref.title.rstrip(".") + ".")
    if ref.journal: bits.append(f"*{ref.journal}*.")
    yvi = ref.year or ""
    if ref.volume: yvi += f";{ref.volume}"
    if ref.issue: yvi += f"({ref.issue})"
    if ref.pages: yvi += f":{ref.pages}"
    if yvi: bits.append(yvi + ".")
    if ref.doi: bits.append(f"doi:{ref.doi}")
    return f"{index}. " + " ".join(bits).strip()


def format_apa(ref: Reference, index: int = 1) -> str:
    """APA 7th ed. — author-date in-text, alphabetic ref list."""
    authors_raw = ref.authors or []
    if len(authors_raw) > 20:
        first19 = ", ".join(_author_last_first(a) for a in authors_raw[:19])
        last = _author_last_first(authors_raw[-1])
        authors = f"{first19}, ... {last}"
    else:
        authors = ", ".join(_author_last_first(a) for a in authors_raw)
    year = f"({ref.year})" if ref.year else "(n.d.)"
    title = ref.title.rstrip(".") + "." if ref.title else ""
    journal = f"*{ref.journal}*" if ref.journal else ""
    vol = f", *{ref.volume}*" if ref.volume else ""
    issue = f"({ref.issue})" if ref.issue else ""
    pages = f", {ref.pages}" if ref.pages else ""
    doi = f". https://doi.org/{ref.doi}" if ref.doi else ""
    return f"{authors} {year}. {title} {journal}{vol}{issue}{pages}{doi}".strip()


def format_harvard(ref: Reference, index: int = 1) -> str:
    """Harvard author-date — common in public health journals."""
    authors = ", ".join(_author_last_first(a) for a in (ref.authors or []))
    year = ref.year or "n.d."
    title = f"'{ref.title.rstrip('.')}'" if ref.title else ""
    journal = f"*{ref.journal}*" if ref.journal else ""
    vi = ""
    if ref.volume: vi = ref.volume
    if ref.issue: vi += f"({ref.issue})"
    pages = f"pp. {ref.pages}" if ref.pages else ""
    doi = f"doi: {ref.doi}" if ref.doi else ""
    return f"{authors} ({year}) {title}, {journal}, {vi}, {pages}. {doi}".strip(" ,.")


def format_ieee(ref: Reference, index: int = 1) -> str:
    """IEEE — biomedical engineering journals."""
    authors = ", ".join(_author_last_first(a) for a in (ref.authors or []))
    title = f'"{ref.title.rstrip(".")},"' if ref.title else ""
    journal = f"*{ref.journal}*," if ref.journal else ""
    vol = f"vol. {ref.volume}," if ref.volume else ""
    no = f"no. {ref.issue}," if ref.issue else ""
    pages = f"pp. {ref.pages}," if ref.pages else ""
    year = ref.year or ""
    return f"[{index}] {authors}, {title} {journal} {vol} {no} {pages} {year}.".strip()


_FORMAT_BY_STYLE = {
    "vancouver": format_vancouver,
    "ama": format_ama,
    "apa": format_apa,
    "harvard": format_harvard,
    "ieee": format_ieee,
    "chicago": format_apa,  # author-date Chicago fallback
}


def format_reference(ref: Reference, index: int = 1, style: str = "vancouver") -> str:
    """저널별 reference style로 단일 ref 포맷.

    style은 journal_registry.JournalStyle.reference_style 값을 그대로 받는다.
    """
    fn = _FORMAT_BY_STYLE.get((style or "vancouver").strip().lower(), format_vancouver)
    return fn(ref, index)


# ── ④ 출력 (참고문헌 목록 / Word / EndNote) ────────────────────────────────
def reference_list_markdown(ordered: List[Reference], style: str = "vancouver") -> str:
    if not ordered:
        return ""
    fn = _FORMAT_BY_STYLE.get((style or "vancouver").strip().lower(), format_vancouver)
    return "\n".join(fn(r, i) for i, r in enumerate(ordered, 1))


def build_cited_docx(title: str, sections: Dict[str, str], ordered: List[Reference],
                       style: str = "vancouver") -> bytes:
    """본문(인용 [n] 또는 (Author, year) 포함) + 참고문헌 목록 → Word 바이트.

    style: vancouver / ama / apa / harvard / ieee / chicago — target journal에 맞춰 선택.
    """
    try:
        from docx import Document
    except ImportError:
        return b""
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    labels = {"abstract": "Abstract", "introduction": "Introduction",
              "methods": "Methods", "results": "Results", "discussion": "Discussion",
              "conclusion": "Conclusion"}
    fn = _FORMAT_BY_STYLE.get((style or "vancouver").strip().lower(), format_vancouver)
    for k in ["abstract", "introduction", "methods", "results", "discussion", "conclusion"]:
        body = (sections.get(k) or "").strip()
        if not body:
            continue
        doc.add_heading(labels[k], level=1)
        doc.add_paragraph(body)
    if ordered:
        doc.add_heading("References", level=1)
        for i, r in enumerate(ordered, 1):
            doc.add_paragraph(fn(r, i))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def endnote_bytes(ordered: List[Reference]) -> bytes:
    return to_endnote_xml(ordered).encode("utf-8") if ordered else b""


def bibtex_bytes(ordered: List[Reference]) -> bytes:
    return to_bibtex(ordered).encode("utf-8") if ordered else b""


# ── ⑤ DOCX 풀라운드트립 — 첨삭(조유선 스타일) + 기존 레퍼런스/넘버링 → EndNote 재구성 ──
_REF_HEADER = re.compile(
    r"(?im)^[\s#*]*(references?|bibliography|참고\s*문헌|문\s*헌|인용\s*문헌)\s*$"
)
_ENTRY_SPLIT = re.compile(r"(?m)^\s*(?:\[(\d+)\]|(\d+)[.)])\s+")
_INTEXT = re.compile(r"\[(\d+(?:\s*[-,]\s*\d+)*)\]")


def extract_docx_text(file_bytes: bytes) -> str:
    """DOCX 바이트 → 텍스트. mammoth(구조/heading 보존) 우선, 실패 시 python-docx 폴백."""
    try:
        import mammoth
        with io.BytesIO(file_bytes) as f:
            return (mammoth.extract_raw_text(f).value or "").strip()
    except Exception as ex:
        _log.warning("mammoth 추출 실패 → python-docx 폴백: %s", ex)
    try:
        from docx import Document
        with io.BytesIO(file_bytes) as f:
            doc = Document(f)
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception as ex:
        _log.warning("python-docx 추출도 실패: %s", ex)
        return ""


def _parse_vancouver_entry(entry: str) -> Reference:
    """참고문헌 1줄(밴쿠버 추정) → Reference. 파싱 실패 시 전체를 title로 보존."""
    entry = " ".join(entry.split()).strip().rstrip(".")
    if not entry:
        return Reference()
    doi = ""
    m = re.search(r"10\.\d{4,9}/\S+", entry)
    if m:
        doi = m.group(0).rstrip(".")
    pmid = ""
    mp = re.search(r"PMID:?\s*(\d{4,9})", entry, re.I)
    if mp:
        pmid = mp.group(1)
    year = ""
    my = re.search(r"\b(19|20)\d{2}\b", entry)
    if my:
        year = my.group(0)
    # 밴쿠버: "Authors. Title. Journal. Year;vol(iss):pages."
    segs = [s.strip() for s in re.split(r"\.\s+", entry) if s.strip()]
    authors: List[str] = []
    title = ""
    journal = ""
    if len(segs) >= 2:
        # 첫 세그가 저자 목록처럼 보이면(콤마+이니셜) 분리
        if re.search(r"[A-Z]{1,3}\s*,|,\s*[A-Z]{1,3}\b|et al", segs[0]) or len(segs[0].split()) <= 12:
            authors = [a.strip() for a in re.split(r",|;", segs[0]) if a.strip()][:30]
            title = segs[1]
            journal = segs[2] if len(segs) >= 3 else ""
        else:
            title = segs[0]
            journal = segs[1] if len(segs) >= 2 else ""
    else:
        title = entry
    journal = re.split(r"\b(19|20)\d{2}\b", journal)[0].strip().rstrip(".,;") if journal else ""
    vol = iss = pages = ""
    mv = re.search(r"(\d+)\s*(?:\((\d+)\))?\s*:\s*([\d\-eE]+)", entry)
    if mv:
        vol, iss, pages = mv.group(1), (mv.group(2) or ""), mv.group(3)
    return Reference(title=title[:400], authors=authors, journal=journal[:200],
                     year=year, volume=vol, issue=iss, pages=pages, doi=doi, pmid=pmid)


def extract_existing_references(full_text: str) -> Dict:
    """DOCX 전문 → {body, refs(번호순 Reference), markers, ref_found}.

    기존 참고문헌 섹션과 본문 [n] 넘버링을 보존해 EndNote 재구성의 입력으로 만든다.
    """
    text = full_text or ""
    m = _REF_HEADER.search(text)
    if not m:
        return {"body": text, "refs": [], "markers": [], "ref_found": False}
    body = text[:m.start()].rstrip()
    ref_block = text[m.end():].strip()
    # 항목 분리: "[1]" 또는 "1." / "1)" 머리번호 기준
    refs: List[Reference] = []
    if _ENTRY_SPLIT.search(ref_block):
        pieces = _ENTRY_SPLIT.split(ref_block)
        # split 결과: [pre, num1a, num1b, entry1, num2a, num2b, entry2, ...]
        buf = pieces[1:]
        for i in range(0, len(buf), 3):
            entry = (buf[i + 2] if i + 2 < len(buf) else "").strip()
            if entry:
                refs.append(_parse_vancouver_entry(entry))
    else:
        for line in ref_block.splitlines():
            if len(line.strip()) > 15:
                refs.append(_parse_vancouver_entry(line))
    markers = sorted({int(n) for grp in _INTEXT.findall(body)
                      for n in re.findall(r"\d+", grp)})
    return {"body": body, "refs": refs, "markers": markers, "ref_found": True}


def restyle_text_yoosun(text: str, style_name: str = "Yoosun Cho",
                        study: Optional[Dict] = None) -> Tuple[str, str]:
    """본문을 조유선 스타일로 첨삭. (rewritten, note) 반환. LLM 실패 시 원문+사유 보존(규칙11)."""
    if not (text or "").strip():
        return text, "빈 본문"
    try:
        from src.llm import get_llm_client
        client = get_llm_client(task="paper_writing")
        sys = (
            f"You are an expert medical writer editing in the academic style of {style_name} "
            "(Korean public-health epidemiology papers: precise, concise, IMRAD, evidence-first). "
            "Proofread and refine the text. PRESERVE all in-text citation markers like [1], [2,3] "
            "EXACTLY and in place. PRESERVE all numbers, statistics, and section headers (## ...). "
            "Do not invent references. Return ONLY the revised full text."
        )
        prompt = (f"STUDY: {study}\n\n원문 첨삭(조유선 스타일):\n\n{text[:12000]}"
                  if study else f"원문 첨삭(조유선 스타일):\n\n{text[:12000]}")
        out = client.generate(prompt, system_prompt=sys, task="paper_writing", max_tokens=4000)
        if out and out.strip() and len(out.strip()) > len(text) * 0.4:
            return out.strip(), "조유선 스타일 첨삭 완료"
        return text, "LLM 응답이 비정상(짧음) — 원문 유지"
    except Exception as ex:
        return text, f"LLM 호출 실패(쿼터/크레딧 가능) — 원문 유지: {str(ex)[:120]}"


def revise_docx_fullset(file_bytes: bytes, style_name: str = "Yoosun Cho",
                        study: Optional[Dict] = None, restyle: bool = True) -> Dict:
    """DOCX 첨삭 풀셋 라운드트립.

    docx → 텍스트(mammoth) → 기존 레퍼런스/넘버링 분리 → 본문 조유선 첨삭(옵션)
    → 번호순 레퍼런스 보존 → 첨삭본 DOCX + EndNote XML + BibTeX.

    반환: {body, refs, n_refs, markers, note, docx, endnote, bibtex, ref_found}
    """
    text = extract_docx_text(file_bytes)
    if not text:
        return {"error": "DOCX에서 텍스트를 추출하지 못했습니다."}
    parsed = extract_existing_references(text)
    body, refs = parsed["body"], parsed["refs"]
    note = "레퍼런스 미검출 — 본문만 첨삭" if not parsed["ref_found"] else f"레퍼런스 {len(refs)}개 보존"
    new_body = body
    if restyle:
        new_body, style_note = restyle_text_yoosun(body, style_name, study)
        note = f"{note} · {style_note}"
    # 첨삭 본문을 섹션으로 재분해(헤더 보존됐다면) → 기존 build_cited_docx 재사용
    sections = _split_body_to_sections(new_body)
    title = sections.pop("_title", "") or (study or {}).get("title", "") or "Revised Paper"
    docx = build_cited_docx(title, sections, refs)
    return {
        "body": new_body, "refs": refs, "n_refs": len(refs),
        "markers": parsed["markers"], "note": note, "ref_found": parsed["ref_found"],
        "docx": docx, "endnote": endnote_bytes(refs), "bibtex": bibtex_bytes(refs),
    }


def _split_body_to_sections(body: str) -> Dict[str, str]:
    """첨삭 본문 → 섹션 dict. '## Label' 또는 표준 IMRAD 헤더 인식, 없으면 introduction 통합."""
    lab2key = {"abstract": "abstract", "introduction": "introduction", "서론": "introduction",
               "methods": "methods", "method": "methods", "방법": "methods",
               "results": "results", "결과": "results",
               "discussion": "discussion", "고찰": "discussion", "논의": "discussion",
               "conclusion": "discussion", "결론": "discussion"}
    parts = re.split(r"(?m)^\s*#{0,3}\s*([A-Za-z가-힣]+)\s*$", body or "")
    head = (parts[0] or "").strip()
    out: Dict[str, str] = {}
    if head:
        out["_title"] = head.splitlines()[0].strip()
        if len(parts) == 1:  # 헤더 전무 → 전체를 introduction
            out["introduction"] = head
            return out
    for i in range(1, len(parts) - 1, 2):
        key = lab2key.get(parts[i].strip().lower())
        seg = (parts[i + 1] or "").strip()
        if key and seg:
            out[key] = (out.get(key, "") + "\n\n" + seg).strip() if out.get(key) else seg
    if not any(k for k in out if k != "_title"):
        out["introduction"] = body.strip()
    return out
