"""Word 논문 출력 엔진 — `data/templates/manuscript_template.json` 양식 1:1 적용.

zcb_dep_v5_0525.pdf의 Yoosun Cho 표준 양식을 그대로 보존하여, 모든 논문 docx 출력이
같은 룩앤필을 가지도록 한다. **template JSON은 단일 진실원본** — 변경 시 version bump.

레이아웃:
  Page 1: Title page (제목/저자/affiliation 중앙)
  Page 2: Abstract + Background/Methods/Results/Conclusion inline bold + Keywords
  Page 3+: Introduction / Methods / Results / Discussion (subsection italic+bold)
  Back: Ethics approval / Data availability / Competing interests / Funding /
        Authors' contributions (italic+bold heading)
  Last: References (Vancouver — title bold, journal italic, volume bold)

호출 예:
    WordExporter().export(
        topic={"title": "...", "authors": ["Yoosun Cho", ...], "affiliations": ["..."]},
        sections={"Abstract": {...}, "Introduction": "...", ...},
        figures=[{"bytes": ..., "caption": "...", "n": 1}],
        tables=[{"type": "baseline", "data": ..., "caption": "...", "n": 1}],
        references=[{"authors": ..., "title": ..., ...}],
        output_path="...",
    )
"""
from __future__ import annotations

import io
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from src.config.logging_config import get_logger

_log = get_logger(__name__)
_OUTPUT_DIR = Path("data/drafts/word")
_TEMPLATE_PATH = Path("data/templates/manuscript_template.json")


# ── Template loading ─────────────────────────────────────────────────────────

_TPL_CACHE: Optional[dict] = None


def load_template(path: Optional[Path] = None) -> dict:
    """단일 진실원본 양식 spec 로드. 캐시됨 — `reload_template()`로 무효화."""
    global _TPL_CACHE
    if _TPL_CACHE is None or path is not None:
        p = path or _TEMPLATE_PATH
        if not p.exists():
            _log.warning("template 없음, fallback: %s", p)
            return _fallback_template()
        _TPL_CACHE = json.loads(p.read_text(encoding="utf-8"))
    return _TPL_CACHE


def reload_template() -> dict:
    global _TPL_CACHE
    _TPL_CACHE = None
    return load_template()


def _fallback_template() -> dict:
    """template 파일이 없을 때 zcb_dep_v5 양식의 최소 spec (안전망)."""
    return {
        "page": {"size": "A4", "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 3.0, "right": 2.5}},
        "font": {"family": "Times New Roman", "body_size_pt": 11, "title_size_pt": 12,
                 "table_size_pt": 9, "caption_size_pt": 10, "reference_size_pt": 11},
        "spacing": {"line_spacing_rule": "double", "first_line_indent_cm": 1.27,
                    "space_before_heading_pt": 12, "space_after_heading_pt": 6},
        "abstract": {"heading_text": "Abstract", "inline_labels":
                     ["Background", "Methods", "Results", "Conclusion"],
                     "label_separator": ": ", "keywords_label": "Keywords",
                     "keyword_separator": "; "},
        "main_sections": {"order": ["Introduction", "Methods", "Results", "Discussion"]},
        "back_matter": {"sections": ["Ethics approval", "Data availability",
                                      "Competing interests", "Funding",
                                      "Authors' contributions"]},
        "references": {"section_heading": "References"},
    }


# ── Style helpers ────────────────────────────────────────────────────────────

def _apply_base_style(doc, tpl: dict):
    """문서 전역 base style — Normal 스타일 + 페이지 마진."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_LINE_SPACING

    font = tpl["font"]
    spacing = tpl["spacing"]
    margins = tpl["page"]["margins_cm"]

    style = doc.styles["Normal"]
    style.font.name = font["family"]
    style.font.size = Pt(font["body_size_pt"])

    # Double spacing 전역 적용
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(spacing.get("paragraph_space_after_pt", 0))

    section = doc.sections[0]
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])


def _set_run(run, *, bold: bool = False, italic: bool = False, size_pt: Optional[int] = None,
             family: Optional[str] = None, superscript: bool = False):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size_pt is not None:
        from docx.shared import Pt
        run.font.size = Pt(size_pt)
    if family is not None:
        run.font.name = family
    if superscript:
        run.font.superscript = True


def _set_paragraph(p, *, alignment: str = "justify", indent_cm: Optional[float] = None,
                   space_before_pt: Optional[int] = None, space_after_pt: Optional[int] = None,
                   line_spacing: str = "double"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Cm, Pt

    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                 "center": WD_ALIGN_PARAGRAPH.CENTER, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    pf = p.paragraph_format
    if indent_cm is not None:
        pf.first_line_indent = Cm(indent_cm)
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if line_spacing == "double":
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif line_spacing == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif line_spacing == "1.5":
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


# ── Section builders ─────────────────────────────────────────────────────────

def _build_title_page(doc, tpl: dict, title: str, authors: List[dict],
                      affiliations: List[str]):
    """zcb_dep_v5 양식 page 1: 제목 중앙 굵게, 저자 중앙(superscript marker), affiliation 중앙."""
    font = tpl["font"]
    tp = tpl["title_page"]

    # 제목
    p = doc.add_paragraph()
    _set_paragraph(p, alignment="center", line_spacing="double",
                   space_before_pt=0, space_after_pt=12)
    run = p.add_run(title)
    _set_run(run, bold=tp.get("title_bold", True),
             size_pt=tp.get("title_size_pt", font["title_size_pt"]),
             family=font["family"])

    # 저자 + superscript marker
    if authors:
        p = doc.add_paragraph()
        _set_paragraph(p, alignment="center", line_spacing="double",
                       space_after_pt=6)
        for i, a in enumerate(authors):
            if i > 0:
                run = p.add_run(", ")
                _set_run(run, size_pt=font["body_size_pt"], family=font["family"])
            name = a.get("name", a) if isinstance(a, dict) else a
            marker = (a.get("affil_marker", "") if isinstance(a, dict)
                      else "")
            bold_first = a.get("first_author", False) if isinstance(a, dict) else False
            run = p.add_run(name)
            _set_run(run, bold=bold_first, size_pt=font["body_size_pt"],
                     family=font["family"])
            if marker:
                run = p.add_run(marker)
                _set_run(run, superscript=True, size_pt=font["body_size_pt"],
                         family=font["family"])

    # Affiliations
    for aff in affiliations:
        p = doc.add_paragraph()
        _set_paragraph(p, alignment="center", line_spacing="single",
                       space_after_pt=2)
        # marker가 앞에 있으면 superscript
        m = re.match(r"^([a-z])(.*)$", aff)
        if m:
            run = p.add_run(m.group(1))
            _set_run(run, superscript=True,
                     size_pt=tp.get("affiliation_size_pt", font["body_size_pt"]),
                     family=font["family"])
            rest = m.group(2).lstrip()
            run = p.add_run(rest)
            _set_run(run, size_pt=tp.get("affiliation_size_pt", font["body_size_pt"]),
                     family=font["family"])
        else:
            run = p.add_run(aff)
            _set_run(run, size_pt=tp.get("affiliation_size_pt", font["body_size_pt"]),
                     family=font["family"])

    if tp.get("separate_page", True):
        doc.add_page_break()


def _build_abstract(doc, tpl: dict, abstract: Any, keywords: List[str]):
    """zcb_dep_v5 양식 page 2: Abstract heading + Background/Methods/Results/Conclusion inline bold."""
    font = tpl["font"]
    ab = tpl["abstract"]
    spacing = tpl["spacing"]

    # Abstract heading
    p = doc.add_paragraph()
    _set_paragraph(p, alignment="left", line_spacing="double",
                   space_before_pt=0, space_after_pt=spacing["space_after_heading_pt"])
    run = p.add_run(ab["heading_text"])
    _set_run(run, bold=ab.get("heading_bold", True),
             size_pt=ab.get("heading_size_pt", font["body_size_pt"]),
             family=font["family"])

    # Abstract body — dict이면 inline label, str이면 그대로
    if isinstance(abstract, dict):
        for label in ab["inline_labels"]:
            content = abstract.get(label) or abstract.get(label.lower())
            if not content:
                continue
            p = doc.add_paragraph()
            _set_paragraph(p, alignment="justify", line_spacing="double",
                           indent_cm=None, space_after_pt=6)
            run = p.add_run(label)
            _set_run(run, bold=ab.get("inline_label_bold", True),
                     size_pt=font["body_size_pt"], family=font["family"])
            run = p.add_run(ab["label_separator"])
            _set_run(run, size_pt=font["body_size_pt"], family=font["family"])
            _add_text_with_inline_cites(p, str(content).strip(), font)
    elif isinstance(abstract, str) and abstract.strip():
        p = doc.add_paragraph()
        _set_paragraph(p, alignment="justify", line_spacing="double")
        _add_text_with_inline_cites(p, abstract.strip(), font)

    # Keywords
    if keywords:
        p = doc.add_paragraph()
        _set_paragraph(p, alignment="justify", line_spacing="double",
                       space_before_pt=6)
        run = p.add_run(ab["keywords_label"])
        _set_run(run, bold=True, size_pt=font["body_size_pt"], family=font["family"])
        run = p.add_run(ab["label_separator"])
        _set_run(run, size_pt=font["body_size_pt"], family=font["family"])
        run = p.add_run(ab["keyword_separator"].join(keywords))
        _set_run(run, size_pt=font["body_size_pt"], family=font["family"])

    doc.add_page_break()


def _build_main_section(doc, tpl: dict, heading: str, body: Any):
    """Introduction / Methods / Results / Discussion. body는 str 또는 subsection dict."""
    font = tpl["font"]
    spacing = tpl["spacing"]
    main = tpl["main_sections"]

    # Section heading
    p = doc.add_paragraph()
    _set_paragraph(p, alignment="left", line_spacing="double",
                   space_before_pt=spacing["space_before_heading_pt"],
                   space_after_pt=spacing["space_after_heading_pt"])
    run = p.add_run(heading)
    _set_run(run, bold=main.get("heading_bold", True),
             italic=main.get("heading_italic", False),
             size_pt=main.get("heading_size_pt", font["body_size_pt"]),
             family=font["family"])

    if isinstance(body, dict):
        # subsection 형태: {"Study population": "...", "Measurements": "...", ...}
        for sub_heading, sub_body in body.items():
            if sub_heading in ("__intro__", "_intro"):
                _add_body_paragraphs(doc, tpl, str(sub_body))
                continue
            # subsection heading (italic + bold)
            p = doc.add_paragraph()
            _set_paragraph(p, alignment="left", line_spacing="double",
                           space_before_pt=8, space_after_pt=4)
            run = p.add_run(sub_heading)
            _set_run(run, bold=main.get("subheading_bold", True),
                     italic=main.get("subheading_italic", True),
                     size_pt=main.get("subheading_size_pt", font["body_size_pt"]),
                     family=font["family"])
            _add_body_paragraphs(doc, tpl, str(sub_body))
    else:
        _add_body_paragraphs(doc, tpl, str(body))


def _add_body_paragraphs(doc, tpl: dict, text: str):
    """본문 단락 추가 — 빈 줄 기준 분리, 첫 줄 들여쓰기, 양쪽 정렬."""
    font = tpl["font"]
    main = tpl["main_sections"]
    indent = main.get("body_first_line_indent_cm",
                       tpl["spacing"].get("first_line_indent_cm", 1.27))

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    for para in paragraphs:
        p = doc.add_paragraph()
        _set_paragraph(p, alignment="justify", line_spacing="double",
                       indent_cm=indent, space_after_pt=0)
        _add_text_with_inline_cites(p, para, font)


_CITE_RE = re.compile(r"\[(\d+(?:[\s,\-–]\d+)*)\]")
_ITALIC_MARKERS = [
    (r"\bP\s*(<|=|>)\s*(0?\.\d+)", "P-value"),       # P < 0.001 → italic P
    (r"\bP\s*for\s*(trend|interaction)\s*=\s*(0?\.\d+)", "P-phrase"),
]
_FIG_TABLE_RE = re.compile(r"\b(Figure|Table|Supplementary Table|Supplementary Figure)\s+(\d+|S?\d+|[A-Z]?\d*)", re.IGNORECASE)


def _add_text_with_inline_cites(p, text: str, font: dict):
    """본문에 [1, 2] / Figure 1 / Table 1 / P < 0.001 등 의학 inline 양식 자동 적용."""
    cursor = 0
    # 합성 토큰: cite (대괄호 그대로 유지 — Vancouver), Figure/Table (Bold), italic P-value
    pattern = re.compile(
        r"(\[\d+(?:[\s,\-–]\d+)*\])"                                       # [1, 2-5]
        r"|((?:Figure|Table|Supplementary\s+(?:Figure|Table))\s+[A-Za-z]?\d+)"  # Figure 1, Table 2, Supplementary Table 1
        r"|(\bP\s*(?:<|=|>)\s*0?\.\d+)"                                      # P < 0.001
        r"|(\bP\s*for\s*(?:trend|interaction)\s*(?:<|=|>)\s*0?\.\d+)"        # P for trend = 0.001
    )
    for m in pattern.finditer(text):
        # 일반 텍스트
        if m.start() > cursor:
            run = p.add_run(text[cursor:m.start()])
            _set_run(run, size_pt=font["body_size_pt"], family=font["family"])
        token = m.group(0)
        cite, ftref, pv1, pv2 = m.group(1), m.group(2), m.group(3), m.group(4)
        if cite:
            run = p.add_run(token)
            _set_run(run, size_pt=font["body_size_pt"], family=font["family"])
        elif ftref:
            run = p.add_run(token)
            _set_run(run, bold=True, size_pt=font["body_size_pt"], family=font["family"])
        elif pv1 or pv2:
            # P만 italic, 숫자는 정상
            sub = re.match(r"(\bP)(\s*(?:for\s*(?:trend|interaction))?\s*(?:<|=|>)\s*0?\.\d+)", token)
            if sub:
                run = p.add_run(sub.group(1))
                _set_run(run, italic=True, size_pt=font["body_size_pt"], family=font["family"])
                run = p.add_run(sub.group(2))
                _set_run(run, size_pt=font["body_size_pt"], family=font["family"])
            else:
                run = p.add_run(token)
                _set_run(run, size_pt=font["body_size_pt"], family=font["family"])
        cursor = m.end()
    if cursor < len(text):
        run = p.add_run(text[cursor:])
        _set_run(run, size_pt=font["body_size_pt"], family=font["family"])


def _build_back_matter(doc, tpl: dict, back_matter: Dict[str, str]):
    """Ethics / Data / Competing / Funding / Authors' contributions (italic+bold heading)."""
    font = tpl["font"]
    bm = tpl["back_matter"]
    spacing = tpl["spacing"]

    for sec in bm["sections"]:
        content = back_matter.get(sec) or back_matter.get(sec.lower().replace(" ", "_"))
        if not content:
            continue
        # heading
        p = doc.add_paragraph()
        _set_paragraph(p, alignment="left", line_spacing="double",
                       space_before_pt=spacing["space_before_heading_pt"],
                       space_after_pt=spacing["space_after_heading_pt"])
        run = p.add_run(sec)
        _set_run(run, bold=bm.get("heading_bold", True),
                 italic=bm.get("heading_italic", True),
                 size_pt=bm.get("heading_size_pt", font["body_size_pt"]),
                 family=font["family"])
        # body
        _add_body_paragraphs(doc, tpl, str(content).strip())


def _build_references(doc, tpl: dict, references: List[Dict]):
    """Vancouver 양식 — title bold, journal italic, volume bold."""
    font = tpl["font"]
    refs = tpl["references"]
    spacing = tpl["spacing"]

    p = doc.add_paragraph()
    _set_paragraph(p, alignment="left", line_spacing="double",
                   space_before_pt=spacing["space_before_heading_pt"],
                   space_after_pt=spacing["space_after_heading_pt"])
    run = p.add_run(refs.get("section_heading", "References"))
    _set_run(run, bold=refs.get("heading_bold", True),
             size_pt=refs.get("heading_size_pt", font["body_size_pt"]),
             family=font["family"])

    for i, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        _set_paragraph(p, alignment="left", line_spacing="single",
                       space_before_pt=0, space_after_pt=4)
        # "1." 번호
        run = p.add_run(f"{i}.")
        _set_run(run, size_pt=refs.get("entry_size_pt", font["reference_size_pt"]),
                 family=font["family"])
        run = p.add_run("\t")
        _set_run(run, size_pt=refs.get("entry_size_pt", font["reference_size_pt"]),
                 family=font["family"])

        if isinstance(ref, dict):
            _render_vancouver_entry(p, ref, font, refs)
        else:
            # 사전 포맷된 문자열
            run = p.add_run(str(ref))
            _set_run(run, size_pt=refs.get("entry_size_pt", font["reference_size_pt"]),
                     family=font["family"])


def _render_vancouver_entry(p, ref: dict, font: dict, refs: dict):
    """단일 reference를 Vancouver 양식 inline runs로 렌더."""
    sz = refs.get("entry_size_pt", font["reference_size_pt"])
    fam = font["family"]

    authors = ref.get("authors", "")
    if isinstance(authors, list):
        authors = ", ".join(authors)
    title = ref.get("title", "")
    journal = ref.get("journal", "")
    year = str(ref.get("year", ""))
    volume = str(ref.get("volume", ""))
    issue = str(ref.get("issue", ""))
    pages = ref.get("pages", "")

    # Authors:
    if authors:
        run = p.add_run(authors)
        _set_run(run, size_pt=sz, family=fam)
        run = p.add_run(": ")
        _set_run(run, size_pt=sz, family=fam)
    # Title (bold)
    if title:
        run = p.add_run(title.rstrip(".") + ".")
        _set_run(run, bold=refs.get("title_bold", True), size_pt=sz, family=fam)
        run = p.add_run(" ")
        _set_run(run, size_pt=sz, family=fam)
    # Journal (italic)
    if journal:
        run = p.add_run(journal)
        _set_run(run, italic=refs.get("journal_italic", True), size_pt=sz, family=fam)
        run = p.add_run(" ")
        _set_run(run, size_pt=sz, family=fam)
    # Year,
    if year:
        run = p.add_run(f"{year}, ")
        _set_run(run, size_pt=sz, family=fam)
    # Volume (bold)
    if volume:
        run = p.add_run(volume)
        _set_run(run, bold=refs.get("volume_bold", True), size_pt=sz, family=fam)
        if issue:
            run = p.add_run(f"({issue})")
            _set_run(run, size_pt=sz, family=fam)
    # :pages.
    if pages:
        run = p.add_run(f":{pages}.")
        _set_run(run, size_pt=sz, family=fam)
    elif volume:
        run = p.add_run(".")
        _set_run(run, size_pt=sz, family=fam)


# ── Figures / Tables embedding ───────────────────────────────────────────────

def _embed_figure(doc, tpl: dict, fig: Dict):
    """그림 embed — 캡션은 'Figure N. {title}.' 양식."""
    from docx.shared import Cm

    f_tpl = tpl["figures"]
    font = tpl["font"]
    img_bytes = fig.get("bytes")
    caption = fig.get("caption", "")
    n = fig.get("n")
    width = fig.get("width_cm", f_tpl.get("width_cm", 14.0))

    if not img_bytes:
        return

    p = doc.add_paragraph()
    _set_paragraph(p, alignment="center", line_spacing="single",
                   space_before_pt=12, space_after_pt=4)
    run = p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Cm(width))

    if caption:
        cp = doc.add_paragraph()
        _set_paragraph(cp, alignment=f_tpl.get("caption_alignment", "left"),
                       line_spacing="single", space_before_pt=2, space_after_pt=12)
        if n is not None and f_tpl.get("caption_bold_label", True):
            label = f_tpl.get("caption_label_pattern", "Figure {n}.").format(n=n)
            run = cp.add_run(label + " ")
            _set_run(run, bold=True, size_pt=f_tpl.get("caption_size_pt", 10),
                     family=font["family"])
        run = cp.add_run(caption)
        _set_run(run, size_pt=f_tpl.get("caption_size_pt", 10), family=font["family"])


def _embed_table(doc, tpl: dict, tbl: Dict):
    """표 embed — table_builder에 위임 (학술지 양식)."""
    try:
        from src.export.table_builder import render_publication_table
    except Exception:
        return
    render_publication_table(doc, tbl, tpl)


# ── Main exporter ────────────────────────────────────────────────────────────

class WordExporter:
    """zcb_dep_v5 양식 1:1 적용 표준 논문 docx 생성기.

    `data/templates/manuscript_template.json`이 단일 진실원본.
    모든 figure/table 임베드는 학술지 수준 양식으로 표준화됨.
    """

    def __init__(self, template_path: Optional[Path] = None):
        self.tpl = load_template(template_path)

    def export(
        self,
        topic: Dict,
        sections: Dict[str, Any],
        figures: Optional[List[Dict]] = None,
        tables: Optional[List[Dict]] = None,
        references: Optional[List[Dict]] = None,
        back_matter: Optional[Dict[str, str]] = None,
        keywords: Optional[List[str]] = None,
        output_path: Optional[str] = None,
    ) -> str:
        """zcb_dep_v5 양식대로 docx 생성.

        Args:
            topic: {"title": str, "authors": [{"name": str, "first_author": bool,
                    "affil_marker": "a"}], "affiliations": ["aDepartment, ...", ...]}
                또는 단순 {"title": str}
            sections: {"Abstract": {"Background": "...", "Methods": "...",
                       "Results": "...", "Conclusion": "..."} or str,
                       "Introduction": str,
                       "Methods": dict (subsection) or str,
                       "Results": dict (subsection) or str,
                       "Discussion": str}
            figures: [{"bytes": bytes, "caption": str, "n": int, "width_cm": float}]
            tables: [{"type": str, "data": ..., "caption": str, "n": int}]
            references: [{"authors": str|list, "title": str, "journal": str,
                          "year": str, "volume": str, "issue": str, "pages": str,
                          "doi": str, "pmid": str}]
            back_matter: {"Ethics approval": str, "Data availability": str, ...}
            keywords: ["zero-calorie beverage", "depression", ...]
            output_path: 저장 경로 (None이면 data/drafts/word/ 자동)

        Returns:
            저장된 .docx 경로
        """
        from docx import Document

        _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        doc = Document()
        _apply_base_style(doc, self.tpl)

        # ── Title page
        title = topic.get("title", "Untitled Manuscript")
        authors = topic.get("authors", [])
        if authors and isinstance(authors[0], str):
            # 단순 문자열 리스트 → dict 변환 (첫번째 first_author)
            authors = [{"name": a, "first_author": (i == 0), "affil_marker": "a"}
                       for i, a in enumerate(authors)]
        affiliations = topic.get("affiliations", [])
        if isinstance(affiliations, str):
            affiliations = [affiliations] if affiliations else []
        _build_title_page(doc, self.tpl, title, authors, affiliations)

        # ── Abstract + Keywords
        abstract = sections.get("Abstract", "")
        if abstract:
            _build_abstract(doc, self.tpl, abstract, keywords or [])

        # ── Main sections
        order = self.tpl["main_sections"]["order"]
        for sec_name in order:
            body = sections.get(sec_name)
            if not body:
                continue
            _build_main_section(doc, self.tpl, sec_name, body)
            # Results 뒤에 figures/tables 삽입
            if sec_name == "Results":
                for tbl in (tables or []):
                    _embed_table(doc, self.tpl, tbl)
                for fig in (figures or []):
                    _embed_figure(doc, self.tpl, fig)

        # ── Back matter
        if back_matter:
            _build_back_matter(doc, self.tpl, back_matter)

        # ── References
        if references:
            _build_references(doc, self.tpl, references)

        # 저장
        if not output_path:
            safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title[:60]).strip()
            output_path = str(_OUTPUT_DIR / f"{safe_title}.docx")

        doc.save(output_path)
        _log.info("Word 저장 완료(zcb_dep_v5 양식): %s", output_path)
        return output_path

    def export_bytes(self, **kwargs) -> bytes:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            self.export(output_path=tmp_path, **kwargs)
            with open(tmp_path, "rb") as f:
                return f.read()
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
